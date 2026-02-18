"""
Data-driven proposal generator. No hardcoded client/address/visits/totals.
Load a project config (JSON), compute Exhibit C totals from row data, fill template and save.
For a new proposal: add a new JSON under project_data/ and run with --config that file.
"""
from datetime import datetime
from pathlib import Path
import json
import os
import sys

BASE_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(BASE_DIR))

try:
    from dotenv import load_dotenv
    load_dotenv(BASE_DIR / ".env")
except ImportError:
    pass

from proposal_generator import (
    get_template_path,
    get_proposal_output_dir,
    sanitize_dirname,
    docx_to_pdf,
    set_all_text_black,
)
from docx import Document


def load_project_config(path: str | Path) -> dict:
    """Load and validate project JSON. Computes est_visits and total_fee from exhibit_c_rows."""
    with open(path, "r", encoding="utf-8") as f:
        config = json.load(f)
    rows = config.get("exhibit_c_rows") or []
    visits = [r.get("visits", 0) for r in rows]
    config["_exhibit_c_visits"] = visits
    config["_est_visits"] = sum(visits)
    config["_total_fee"] = config["_est_visits"] * config.get("price_per_visit", 350)
    scope = (config.get("exhibit_a_scope_only") or "").strip()
    addr = (config.get("project_address") or "").strip()
    client = (config.get("client_short") or "").strip()
    config["_exhibit_a_full"] = (
        f"{scope} The project address is {addr}. "
        f"BCC's role on the project will be to serve as the combo inspection inspector, assisting {client} with all required inspections."
    ).strip()
    config["_date"] = datetime.now().strftime("%m-%d-%Y")
    config["_date_long"] = datetime.now().strftime("%B %d, %Y")
    return config


def build_replacements(config: dict) -> dict:
    """Build placeholder -> value map from config. Totals come from computed _est_visits and _total_fee.
    Includes template leftovers (Cox & Company, Bryan Kerr, 701 Monroe, Insomnia Cookies) so they are always replaced.
    """
    c = config
    ppv = c.get("price_per_visit", 350)
    client_short = c.get("client_short", "")
    attention = c.get("attention", "")
    project_name = c.get("project_name", "")
    project_address = c.get("project_address", "")
    return {
        "{{Client}}": client_short,
        "{{ClientFull}}": c.get("client_full", ""),
        "{{Project}}": project_name,
        "{{ProjectName}}": project_name,
        "{{ProjectAddress}}": project_address,
        "{{Address}}": project_address,
        "{{Attention}}": attention,
        "{{ClientEmail}}": c.get("client_email", ""),
        "{{PricePerVisit}}": str(ppv),
        "{{EstVisits}}": str(c["_est_visits"]),
        "{{Total}}": str(c["_total_fee"]),
        "{{ScopeNotes}}": c["_exhibit_a_full"],
        "{{ExhibitA_Scope}}": c["_exhibit_a_full"],
        "{{ProjectDescription}}": c["_exhibit_a_full"],
        "{{BuildingCodeConsulting}}": "Building Code Consulting LLC",
        "{{DC Third-Party}}": "DC Third-Party agency",
        "{{Yue Cao}}": "Yue Cao (PE, MCP)",
        "{{Date}}": c["_date"],
        "{{DateLong}}": c["_date_long"],
        "{{ProjectSize}}": c.get("project_size_sqft", ""),
        "$325": f"${ppv}",
        "325": str(ppv),
        "Flat rate of $325/visit": f"Flat rate of ${ppv}/visit",
        "$1,300": f"${c['_total_fee']:,}",
        "$1300": f"${c['_total_fee']:,}",
        # Template leftovers: replace in Scope of Work and everywhere else
        "Cox & Company, LLC": client_short,
        "Cox & Company": client_short,
        "Bryan Kerr": attention,
        "701 Monroe Street NE": project_address,
        "701 Monroe ST NE": project_address,
        "701 Monroe": project_address,
        "Monroe Street NE": project_address,
        "Insomnia Cookies Renovation": project_name,
        "Insomnia Cookies": project_name,
        "01-12-2026": c["_date"],
        "01/12/2026": c["_date"],
    }


def replace_in_paragraph(para, replacements: dict, config: dict):
    for run in para.runs:
        for k, v in replacements.items():
            if k in run.text:
                run.text = run.text.replace(k, str(v))
        addr = config.get("project_address", "")
        if "United States of America" in run.text:
            run.text = run.text.replace(", United States of America", "").replace("United States of America", "")
        if run.text.strip().endswith("Washington DC") and addr and "20002" in run.text:
            run.text = run.text.replace(" Washington DC", "").rstrip()


def apply_exhibit_c_table(doc, config: dict):
    """Fill Exhibit C body rows from config.exhibit_c_rows; Total row from computed _est_visits and _total_fee."""
    keywords = [r.get("keyword", "").lower() for r in config.get("exhibit_c_rows", [])]
    visits_list = config["_exhibit_c_visits"]
    price = config.get("price_per_visit", 350)
    est_visits = config["_est_visits"]
    total_fee = config["_total_fee"]
    visits_col, total_col = 1, 3

    for table in doc.tables:
        table_text = " ".join(cell.text for row in table.rows for cell in row.cells).lower()
        if "total" not in table_text or "visit" not in table_text:
            continue
        if len(table.rows[0].cells) < 4:
            continue
        for row in table.rows:
            first_cell = (row.cells[0].text or "").lower()
            row_any = " ".join(c.text or "" for c in row.cells).lower()
            v_cell = row.cells[visits_col].text.strip() if visits_col < len(row.cells) else ""
            t_cell = (row.cells[total_col].text or "").strip() if total_col < len(row.cells) else ""
            is_total_row = (
                "total" in row_any
                and len(row.cells) > total_col
                and (v_cell.isdigit() or "$" in t_cell)
                and "price per" not in row_any
            )
            if is_total_row:
                for ci, cell in enumerate(row.cells):
                    for p in cell.paragraphs:
                        if ci == visits_col:
                            p.clear()
                            p.add_run(str(est_visits))
                        elif ci == total_col:
                            p.clear()
                            p.add_run(f"${total_fee:,}")
                continue
            if "total" in first_cell:
                continue
            for ki, kw in enumerate(keywords):
                if kw and kw in first_cell and ki < len(visits_list):
                    v = visits_list[ki]
                    row_total = v * price
                    for p in row.cells[visits_col].paragraphs:
                        p.clear()
                        p.add_run(str(v))
                    for p in row.cells[total_col].paragraphs:
                        p.clear()
                        p.add_run(f"${row_total:,}")
                    break
        for cell in (c for row in table.rows for c in row.cells):
            for p in cell.paragraphs:
                for run in p.runs:
                    if "$325" in run.text:
                        run.text = run.text.replace("$325", f"${price}")
                    if run.text.strip() == "325":
                        run.text = str(price)
        for cell in (c for row in table.rows for c in row.cells):
            if cell.text.strip() == "$325":
                for p in cell.paragraphs:
                    p.clear()
                    p.add_run(f"${price}")


def generate_proposal(config_path: str | Path) -> Path:
    """Load config, fill template, apply Exhibit C from data (totals computed), save and return docx path."""
    config = load_project_config(config_path)
    template_path = get_template_path("code_compliance")
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")
    out_dir = get_proposal_output_dir(config["client_short"], config["project_name"])
    docx_name = f"{sanitize_dirname(config['project_name'])} - Third Party Code Inspection Proposal from BCC.docx"
    out_docx = out_dir / docx_name

    doc = Document(str(template_path))
    replacements = build_replacements(config)

    for p in doc.paragraphs:
        replace_in_paragraph(p, replacements, config)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, replacements, config)

    # Second pass: paragraph-level replace so we catch template text split across runs (e.g. Scope of Work "Cox & Company, LLC")
    for p in doc.paragraphs:
        full = p.text
        changed = False
        for find_str, repl in replacements.items():
            if find_str in full:
                full = full.replace(find_str, str(repl))
                changed = True
        if changed:
            p.clear()
            p.add_run(full)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full = p.text
                    changed = False
                    for find_str, repl in replacements.items():
                        if find_str in full:
                            full = full.replace(find_str, str(repl))
                            changed = True
                    if changed:
                        p.clear()
                        p.add_run(full)

    for p in doc.paragraphs:
        full = p.text
        if "tenant fit out" in full and "AHUs" in full:
            p.clear()
            p.add_run(config["_exhibit_a_full"])
        elif full.strip() == "Washington DC":
            p.clear()
        elif "01-12-2026" in full or full.strip() == "01/12/2026":
            p.clear()
            p.add_run(config["_date"])
        elif "Flat rate of $325" in full or "(Flat rate of $325" in full:
            p.clear()
            p.add_run(f"Inspection Services Estimated: (Flat rate of ${config['price_per_visit']}/visit)")

    apply_exhibit_c_table(doc, config)
    set_all_text_black(doc)

    try:
        doc.save(str(out_docx))
    except PermissionError:
        out_docx = out_dir / (out_docx.stem + " - CORRECTED.docx")
        doc.save(str(out_docx))
        print("(Original file was open; saved as CORRECTED copy.)")
    return out_docx


def main():
    import argparse
    ap = argparse.ArgumentParser(description="Generate proposal from project config JSON (data-driven, no hardcoded totals).")
    ap.add_argument("config", nargs="?", default=None, help="Path to project_data/*.json or project_id")
    args = ap.parse_args()
    if args.config is None:
        args.config = BASE_DIR / "project_data" / "st_josephs_capitol_hill_phase1.json"
    else:
        p = Path(args.config)
        if not p.suffix:
            args.config = BASE_DIR / "project_data" / f"{args.config}.json"
        elif not p.is_absolute():
            args.config = BASE_DIR / args.config
    path = generate_proposal(args.config)
    print("Proposal saved:", path)
    pdf_path = docx_to_pdf(path)
    if pdf_path:
        print("PDF saved:", pdf_path)
    print("Output path:", path.parent.resolve())


if __name__ == "__main__":
    main()
