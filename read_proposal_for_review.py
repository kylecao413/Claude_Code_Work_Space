"""Extract text from generated proposal so the agent can read and self-approve before sending to user."""
import sys
from pathlib import Path
from docx import Document

BASE = Path(__file__).resolve().parent
# Output dir from proposal generator (same as regenerate_st_josephs_proposal / proposal_from_config)
out_dir = Path(r"c:\Users\Kyle Cao\DC Business\Building Code Consulting\Projects\Keller Brothers\St. Joseph's on Capitol Hill – Phase I")
for name in [
    "St. Joseph's on Capitol Hill – Phase I - Third Party Code Inspection Proposal from BCC - CORRECTED.docx",
    "St. Joseph's on Capitol Hill – Phase I - Third Party Code Inspection Proposal from BCC.docx",
]:
    p = out_dir / name
    if p.exists():
        doc = Document(str(p))
        print("=== PARAGRAPHS ===")
        for i, para in enumerate(doc.paragraphs):
            t = para.text.strip()
            if t:
                print(f"P{i}: {t[:300]}" + ("..." if len(t) > 300 else ""))
        print("\n=== TABLES ===")
        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                row_text = " | ".join(c.text.strip() for c in row.cells)
                if row_text.strip():
                    print(f"T{ti}R{ri}: {row_text[:350]}")
        sys.exit(0)
print("No proposal docx found.", file=sys.stderr)
sys.exit(1)
