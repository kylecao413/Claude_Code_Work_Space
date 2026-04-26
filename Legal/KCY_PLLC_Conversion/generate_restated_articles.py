"""Generate Restated Articles of Organization for KCY Engineer PLLC.

Output: Restated_Articles_of_Organization_KCY_Engineer_PLLC.docx
Kyle: open in Word -> File -> Save As -> PDF -> upload to VA SCC.
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from pathlib import Path

OUT = Path(__file__).parent / "Restated_Articles_of_Organization_KCY_Engineer_PLLC.docx"
SIG = Path(r"C:\Users\Kyle Cao\DC Business\Building Code Consulting\Logo E-Sig Stamp\E-Sig.jpg")

doc = Document()

# Page margins
for section in doc.sections:
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

# Default font
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)


def add_centered(text, bold=False, size=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    return p


def add_heading_article(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    return p


def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


# Title — filing entity must be identified by its CURRENT legal name on file with VA SCC.
# The new name (KCY Engineer PLLC) is adopted by the restatement itself in Article I.
add_centered("RESTATED ARTICLES OF ORGANIZATION", bold=True, size=14)
add_centered("OF", bold=True, size=12)
add_centered("KCY ENGINEERING CODE CONSULTING, LLC", bold=True, size=14)
add_centered("Virginia SCC Entity ID: 11579383", size=11)
add_centered("Originally formed August 10, 2023", size=11)

doc.add_paragraph()  # spacer

add_body(
    "Pursuant to Virginia Code § 13.1-1014.1, the Articles of Organization "
    "of KCY Engineering Code Consulting, LLC (the \"Company\"), originally "
    "filed with the Virginia State Corporation Commission on August 10, 2023, "
    "are hereby restated in their entirety to read as follows:"
)

# ARTICLE I — NAME
add_heading_article("ARTICLE I — NAME")
add_body("The name of the Company is KCY Engineer PLLC.")

# ARTICLE II — ENTITY TYPE
add_heading_article("ARTICLE II — ENTITY TYPE AND STATUTORY FRAMEWORK")
add_body(
    "The Company is a Professional Limited Liability Company organized "
    "under the laws of the Commonwealth of Virginia pursuant to "
    "Title 13.1, Chapter 13 of the Code of Virginia "
    "(§§ 13.1-1100 through 13.1-1123), together with the provisions of "
    "Chapter 12 (§§ 13.1-1000 et seq.) made applicable to professional "
    "limited liability companies by § 13.1-1122."
)

# ARTICLE III — PROFESSIONAL SERVICE
add_heading_article("ARTICLE III — PROFESSIONAL SERVICE")
add_body(
    "The sole professional service for which the Company is organized to "
    "render is the practice of professional engineering as regulated by "
    "the Virginia Board for Architects, Professional Engineers, Land "
    "Surveyors, Certified Interior Designers and Landscape Architects "
    "(APELSCIDLA Board) under Title 54.1, Chapter 4 of the Code of Virginia."
)

# ARTICLE IV — MEMBERS AND MANAGERS
add_heading_article("ARTICLE IV — MEMBERS AND MANAGERS")
add_body(
    "All members and managers of the Company are, and at all times shall "
    "be, duly licensed or otherwise legally authorized to render professional "
    "engineering services in the Commonwealth of Virginia, in accordance "
    "with § 13.1-1102 of the Code of Virginia."
)
add_body(
    "The sole member of the Company is Yue Cao, who holds an active "
    "Virginia Professional Engineer license issued by the APELSCIDLA Board."
)

# ARTICLE V — PRINCIPAL OFFICE
add_heading_article("ARTICLE V — PRINCIPAL OFFICE")
add_body(
    "The address of the principal office of the Company is:"
)
add_body("3914 Tallow Tree Court, Fairfax, Virginia 22033")

# ARTICLE VI — REGISTERED AGENT
add_heading_article("ARTICLE VI — REGISTERED AGENT")
add_body(
    "The name and Virginia business office address of the Company's "
    "registered agent is:"
)
add_body("Yue Cao")
add_body("3914 Tallow Tree Court, Fairfax, Virginia 22033")
add_body(
    "The registered agent is a member of the Company and is a resident "
    "of the Commonwealth of Virginia, qualifying as registered agent "
    "under § 13.1-1015 of the Code of Virginia."
)

# ARTICLE VII — DURATION
add_heading_article("ARTICLE VII — DURATION")
add_body("The duration of the Company is perpetual.")

# ARTICLE VIII — EFFECTIVE DATE
add_heading_article("ARTICLE VIII — EFFECTIVE DATE")
add_body(
    "These Restated Articles of Organization shall become effective upon "
    "filing by the Virginia State Corporation Commission."
)

# Adoption block
doc.add_paragraph()  # spacer
add_heading_article("ADOPTION")
add_body(
    "These Restated Articles of Organization were duly adopted by the "
    "sole member of the Company in accordance with the Virginia Limited "
    "Liability Company Act on April 20, 2026."
)

# Signature block
doc.add_paragraph()
doc.add_paragraph()

# Insert signature image above the signature line
if SIG.exists():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(SIG), width=Inches(1.6))

p = doc.add_paragraph("_______________________________________")
p.paragraph_format.space_after = Pt(0)
p = doc.add_paragraph("Yue Cao, Sole Member")
p.paragraph_format.space_after = Pt(0)
p = doc.add_paragraph("KCY Engineering Code Consulting, LLC")
p.paragraph_format.space_after = Pt(0)
p = doc.add_paragraph("(to be renamed KCY Engineer PLLC upon effectiveness of this filing)")
p.paragraph_format.space_after = Pt(0)
doc.add_paragraph("Date: April 20, 2026")

doc.save(OUT)
print(f"Generated: {OUT}")
