"""Generate Articles of Restatement (cover document) for KCY Engineer PLLC.

Per VA SCC rejection 2026-04-28: a separately typed Articles of Restatement
is required under § 13.1-1014.1. The Restated Articles of Organization
(Restated_Articles_of_Organization_KCY_Engineer_PLLC.pdf) is incorporated by
reference as an attachment.

Output: Articles_of_Restatement_KCY_Engineer_PLLC.docx
Kyle: open in Word -> File -> Save As -> PDF -> upload to VA SCC alongside
the Restated Articles of Organization PDF.
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from pathlib import Path

OUT = Path(__file__).parent / "Articles_of_Restatement_KCY_Engineer_PLLC.docx"
SIG = Path(r"C:\Users\Kyle Cao\DC Business\Building Code Consulting\Logo E-Sig Stamp\E-Sig.jpg")

doc = Document()

for section in doc.sections:
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)


def add_centered(text, bold=False, size=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    return p


def add_heading_article(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    return p


def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


# Title — uses CURRENT legal name on file with VA SCC.
add_centered("ARTICLES OF RESTATEMENT", bold=True, size=14)
add_centered("OF", bold=True, size=12)
add_centered("KCY ENGINEERING CODE CONSULTING, LLC", bold=True, size=14)
add_centered("Virginia SCC Entity ID: 11579383", size=11)

doc.add_paragraph()  # spacer

add_body(
    "Pursuant to § 13.1-1014.1 of the Code of Virginia, KCY Engineering "
    "Code Consulting, LLC, a Virginia limited liability company (the "
    "\"Company\"), hereby adopts the following Articles of Restatement:"
)

# ARTICLE I — CURRENT NAME
add_heading_article("ARTICLE I — CURRENT NAME OF THE COMPANY")
add_body(
    "The current name of the Company on file with the Virginia State "
    "Corporation Commission is KCY Engineering Code Consulting, LLC."
)

# ARTICLE II — FORMATION
add_heading_article("ARTICLE II — DATE OF ORGANIZATION")
add_body(
    "The Articles of Organization of the Company were originally filed "
    "with the Virginia State Corporation Commission on August 10, 2023, "
    "under Virginia SCC Entity ID 11579383."
)

# ARTICLE III — RESTATEMENT
add_heading_article("ARTICLE III — RESTATEMENT")
add_body(
    "The Articles of Organization of the Company are hereby amended and "
    "restated in their entirety as set forth in the Restated Articles of "
    "Organization attached hereto and incorporated herein by reference."
)

# ARTICLE IV — AMENDMENT INCLUDED IN RESTATEMENT
add_heading_article("ARTICLE IV — AMENDMENT CONTAINED IN THE RESTATEMENT")
add_body(
    "The Restated Articles of Organization contain an amendment to the "
    "Articles of Organization, namely the change of the Company's name "
    "from KCY Engineering Code Consulting, LLC to KCY Engineer PLLC, "
    "and the conversion of the Company's status to a Professional Limited "
    "Liability Company organized to render professional engineering "
    "services pursuant to Title 13.1, Chapter 13 of the Code of Virginia."
)

# ARTICLE V — ADOPTION
add_heading_article("ARTICLE V — ADOPTION")
add_body(
    "These Articles of Restatement, together with the attached Restated "
    "Articles of Organization, were duly adopted by the sole member of "
    "the Company in accordance with the Virginia Limited Liability "
    "Company Act on April 20, 2026."
)

# ARTICLE VI — EFFECTIVE DATE
add_heading_article("ARTICLE VI — EFFECTIVE DATE")
add_body(
    "These Articles of Restatement shall become effective upon filing "
    "by the Virginia State Corporation Commission."
)

# Signature block
doc.add_paragraph()
doc.add_paragraph()

if SIG.exists():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(SIG), width=Inches(1.6))

p = doc.add_paragraph("_______________________________________")
p.paragraph_format.space_after = Pt(0)
p = doc.add_paragraph("Yue Cao, Sole Member")
p.paragraph_format.space_after = Pt(0)
p = doc.add_paragraph("KCY Engineering Code Consulting, LLC")
p.paragraph_format.space_after = Pt(0)
doc.add_paragraph("Date: April 20, 2026")

doc.save(OUT)
print(f"Generated: {OUT}")
