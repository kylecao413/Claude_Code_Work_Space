"""Generate the 5-doc DC DOB Plan Review Wrap-Up package for 1522 Rhode Island Ave NE (FA + SPK).

Corrected workflow per Kyle 2026-04-28:
  Doc #1 — copy approved-NOI PDF as-is (preserve filename) into Wrap up/.
  Doc #2 — drive Excel to print the deficiency-report xlsx used range to PDF (one per discipline).
  Doc #3 — fill the AcroForm template on page 61 of the Third-Party_Program_Procedure_Manual,
           extract that single page only, save into Wrap up/. Do NOT recreate in Word.
  Doc #4 — copy/edit the certification-letter .docx; list ALL three disciplines explicitly
           (Fire Protection System / Sprinkler System / Fire Alarm System) — DOB officers
           need it spelled out.
  Doc #5 — copy/edit Large Stamp .docx; only update DATE cells. Never touch signature column —
           the template already contains Yue's signature image embedded there.
"""
from __future__ import annotations
import shutil
from copy import deepcopy
from pathlib import Path
import fitz
from docx import Document
from docx.oxml.ns import qn

# ---------- Project parameters ----------
PROJECT_ROOT = Path(r"C:\Users\Kyle Cao\DC Business\Building Code Consulting\Projects\Nery Soto Fire Protection\1522 Rhode Island Ave NE - PR")
WRAPUP_DIR = PROJECT_ROOT / "Wrap up"
ADDR_SHORT = "1522 Rhode Island Ave NE"
ADDR_FULL = "1522 Rhode Island Ave NE, Washington DC"
ADDR_FULL_DOC5 = "1522 Rhode Island Ave NE, Washington D.C."
PROJECT_NAME = "1522 Condominium LLC"
TPR_NUMBER = "TPR200000-619"
DISCIPLINE_FILE_TAG = "FP"
DEFICIENCY_DATE = "04-21-2026"   # DOB official acceptance date (FOR-OFFICIAL-USE box on approved NOI)
VERIFIED_DATE = "04-28-2026"     # today
APPROVAL_DATE = "04-28-2026"     # today; = signature date on Doc #3
SIGN_DAYMONTH = "April 28th"     # widget value before literal "day of"
SIGN_YEAR_2DIGIT = "26"
LETTER_ISSUING = "04-28-2026"
DOC3_DISC_ROWS = [
    ("Fire Protection",  DEFICIENCY_DATE, VERIFIED_DATE, APPROVAL_DATE),
    ("Sprinkler System", DEFICIENCY_DATE, VERIFIED_DATE, APPROVAL_DATE),
]
DISC_BLANK_DOC3 = "Fire(electrical) & Sprinkler"
LETTER_DISCIPLINES = ["Fire Protection System", "Sprinkler System", "Fire Alarm System"]
DOC5_DATE = APPROVAL_DATE.replace("-", "/")  # mm/dd/yyyy in stamp template

# ---------- Source templates ----------
TPL_DIR = Path(r"C:\Users\Kyle Cao\DC Business\Building Code Consulting\PLAN REVIEW PROJECT DOC TEMPLATE")
GUIDE_MANUAL = TPL_DIR / "Third-Party_Program_Procedure_Manual 5.15.2023 seal.pdf"
APPROVAL_CERT_PAGE = 60  # 0-indexed page 61 of guide manual

WRAP_TPL_DIR = Path(r"C:\Users\Kyle Cao\DC Business\Building Code Consulting\Business Automation\Wrap up")
APPROVED_NOI = PROJECT_ROOT / "Approved_1522_Rhode_Island_Ave_NE_SPK__FA_PR_NOI_signed.pdf"
CERT_LETTER_TPL = WRAP_TPL_DIR / "1345 Connecticut Ave NW Fire Alarm Plan Review Certification Letter.docx"
LARGE_STAMP_TPL = WRAP_TPL_DIR / "Large Stamp Template.docx"

WRAPUP_DIR.mkdir(parents=True, exist_ok=True)


# ---------- Helpers ----------
def _replace_first_text_in_paragraph(p, old: str, new: str) -> bool:
    """Replace `old` with `new` in a paragraph by editing the first run that
    contains the substring; preserves runs that hold images / formatting."""
    if old not in p.text:
        return False
    # Find first run with `old` substring; collapse if split across runs by
    # fusing into the first run and clearing the rest.
    runs_text = "".join(r.text for r in p.runs)
    if old in runs_text and p.runs:
        new_text = runs_text.replace(old, new)
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""
        return True
    return False


def _replace_text_runs_only(d: Document, mapping: dict) -> None:
    """Find/replace across paragraphs and tables WITHOUT clearing cells —
    preserves embedded images (signatures) in the Large Stamp template."""
    def visit(p):
        for old, new in mapping.items():
            _replace_first_text_in_paragraph(p, old, new)

    for p in d.paragraphs:
        visit(p)
    for tab in d.tables:
        for row in tab.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    visit(p)
                for ntab in cell.tables:
                    for nrow in ntab.rows:
                        for ncell in nrow.cells:
                            for np_ in ncell.paragraphs:
                                visit(np_)


# ---------- Doc #1: Copy approved NOI as-is ----------
def doc1_noi() -> Path:
    out = WRAPUP_DIR / APPROVED_NOI.name
    shutil.copy2(APPROVED_NOI, out)
    return out


# ---------- Doc #2: Drive Excel to export deficiency-report xlsx → PDF ----------
def _xlsx_to_pdf(xlsx_path: Path, pdf_out: Path) -> None:
    import win32com.client as win32
    excel = win32.gencache.EnsureDispatch("Excel.Application")
    excel.Visible = False
    excel.DisplayAlerts = False
    try:
        wb = excel.Workbooks.Open(str(xlsx_path), ReadOnly=True)
        try:
            for ws in wb.Worksheets:
                used = ws.UsedRange
                # Force print area to used range and fit-to-1-page-wide
                ws.PageSetup.PrintArea = used.Address
                ws.PageSetup.Zoom = False
                ws.PageSetup.FitToPagesWide = 1
                ws.PageSetup.FitToPagesTall = False
                ws.PageSetup.Orientation = 2  # xlLandscape
            # 0 = xlTypePDF
            wb.ExportAsFixedFormat(0, str(pdf_out))
        finally:
            wb.Close(SaveChanges=False)
    finally:
        excel.Quit()


def doc2_deficiency_pdfs() -> list[Path]:
    spk_xlsx = PROJECT_ROOT / "SPK" / "1522 Rhode Island Ave NE SPK PR Deficiency Report.xlsx"
    fa_xlsx = PROJECT_ROOT / "FA" / "1522 Rhode Island Ave NE FA PR Deficiency Report.xlsx"
    spk_pdf = WRAPUP_DIR / f"{ADDR_SHORT} SPK Deficiency Report.pdf"
    fa_pdf = WRAPUP_DIR / f"{ADDR_SHORT} FA Deficiency Report.pdf"
    out = []
    for src, dst in [(spk_xlsx, spk_pdf), (fa_xlsx, fa_pdf)]:
        try:
            _xlsx_to_pdf(src, dst)
            out.append(dst)
        except Exception as e:
            print(f"   !! xlsx→PDF failed for {src.name}: {e}")
    return out


# ---------- Doc #3: Fill AcroForm on page 61 of guide manual, save single page ----------
def doc3_approval_cert_pdf() -> Path:
    """Open the Third-Party_Program_Procedure_Manual, copy ONLY page 61 to a new
    PDF, fill AcroForm widgets, and save into the project Wrap up folder.
    Yue applies signature to SIGNATURE_3 widget afterwards."""
    out = WRAPUP_DIR / f"{ADDR_SHORT} {DISCIPLINE_FILE_TAG} Plan Review Approval Certificate and Report.pdf"

    src = fitz.open(GUIDE_MANUAL)
    new = fitz.open()
    new.insert_pdf(src, from_page=APPROVAL_CERT_PAGE, to_page=APPROVAL_CERT_PAGE)
    src.close()

    page = new[0]
    # Map field name → value
    field_values = {
        "fill_1": TPR_NUMBER,
        "Date": APPROVAL_DATE,
        "Permit Number": "",
        "Project Name_2": PROJECT_NAME,
        "Project Address_2": ADDR_FULL,
        # Discipline rows
        "Plan Review DisciplineRow1": DOC3_DISC_ROWS[0][0],
        "fill_15": DOC3_DISC_ROWS[0][1],
        "fill_16": DOC3_DISC_ROWS[0][2],
        "fill_17": DOC3_DISC_ROWS[0][3],
        "Plan Review DisciplineRow2": DOC3_DISC_ROWS[1][0],
        "fill_19": DOC3_DISC_ROWS[1][1],
        "fill_20": DOC3_DISC_ROWS[1][2],
        "fill_21": DOC3_DISC_ROWS[1][3],
        # Approval block
        "day of": SIGN_DAYMONTH,
        "20": SIGN_YEAR_2DIGIT,
        "Print Full Name and Title": "Yue Cao Professional in Charge",
        "ProfessionalinCharge of Third Party Plan Review Agency for": DISC_BLANK_DOC3,
        "Name of Agency": "Building Code Consulting",
        "Agency Approval ID Number": "TPR-05012025",
        "Professional EngineerArchitect or MCP Number": "PE920502",
    }
    checkbox_on = {"Fire_2"}

    for w in page.widgets() or []:
        name = w.field_name
        if name in field_values:
            w.field_value = field_values[name]
            w.update()
        elif name in checkbox_on:
            w.field_value = True
            w.update()

    new.save(out, garbage=4, deflate=True)
    new.close()
    return out


# ---------- Doc #4: Certification Letter — list 3 disciplines explicitly ----------
def doc4_cert_letter() -> Path:
    out = WRAPUP_DIR / f"{ADDR_SHORT} Plan Review Certification Letter.docx"
    shutil.copy2(CERT_LETTER_TPL, out)
    d = Document(out)
    mapping = {
        "12-09-2025": LETTER_ISSUING,
        "1345 Connecticut Ave NW Washington DC": ADDR_FULL,
        "Fire Protection System": "::DISC_PLACEHOLDER::",
    }
    _replace_text_runs_only(d, mapping)

    for p in list(d.paragraphs):
        if p.text.strip() == "::DISC_PLACEHOLDER::":
            # Set first paragraph to first discipline
            p.runs[0].text = LETTER_DISCIPLINES[0]
            for r in p.runs[1:]:
                r.text = ""
            xml_parent = p._p.getparent()
            idx = list(xml_parent).index(p._p)
            for j, disc in enumerate(LETTER_DISCIPLINES[1:], start=1):
                new_p = deepcopy(p._p)
                xml_parent.insert(idx + j, new_p)
                ts = new_p.findall('.//' + qn('w:t'))
                if ts:
                    ts[0].text = disc
                    for t in ts[1:]:
                        t.text = ""
            break
    d.save(out)
    return out


# ---------- Doc #5: Large Stamp — preserve embedded signature images ----------
def doc5_large_stamp() -> Path:
    out = WRAPUP_DIR / f"Large Stamp for {ADDR_SHORT} {DISCIPLINE_FILE_TAG} review.docx"
    shutil.copy2(LARGE_STAMP_TPL, out)
    d = Document(out)

    # Replace placeholder address line in outer cell paragraphs (no images there)
    _replace_text_runs_only(d, {
        "1812 H PL SE, Washington D.C.": ADDR_FULL_DOC5,
    })

    # Update DATE cells in nested table — only column 2 (date), never touch
    # column 1 (signature image) or column 0 (label, already correct).
    # Old date may be split across multiple runs (e.g. "11/25" + "/2025"); rewrite
    # the first paragraph cleanly without disturbing any image runs.
    import re
    date_re = re.compile(r"\d{1,2}[/-]\d{1,2}[/-]\d{2,4}")
    DRAWING_TAG = qn("w:drawing")
    PICT_TAG = qn("w:pict")
    for tab in d.tables:
        for row in tab.rows:
            for cell in row.cells:
                for ntab in cell.tables:
                    for nrow in ntab.rows:
                        date_cell = nrow.cells[2]
                        for p in date_cell.paragraphs:
                            full = "".join(r.text for r in p.runs)
                            if not date_re.search(full):
                                continue
                            new_text = date_re.sub(DOC5_DATE, full, count=1)
                            # A run holds an image if its element contains
                            # <w:drawing> or <w:pict> as descendants (not just
                            # the namespace declaration).
                            placed = False
                            for r in p.runs:
                                has_img = (r._r.find(DRAWING_TAG) is not None
                                           or r._r.find(PICT_TAG) is not None)
                                if has_img:
                                    continue
                                if not placed:
                                    r.text = new_text
                                    placed = True
                                else:
                                    r.text = ""
                            break
    d.save(out)
    return out


def main():
    print(f"Generating wrap-up package into: {WRAPUP_DIR}")
    p1 = doc1_noi();           print(f"  [1] {p1.name}")
    p2 = doc2_deficiency_pdfs()
    for p in p2:               print(f"  [2] {p.name}")
    p3 = doc3_approval_cert_pdf(); print(f"  [3] {p3.name}")
    p4 = doc4_cert_letter();   print(f"  [4] {p4.name}")
    p5 = doc5_large_stamp();   print(f"  [5] {p5.name}")
    print("\nManual steps remaining:")
    print("   - Open Doc #3 PDF, sign on the SIGNATURE_3 line, Save.")
    print("   - Open Doc #4 .docx, Save As PDF (auto-signature on letterhead).")
    print("   - Open Doc #5 .docx (signature image already embedded), Save As PDF.")


if __name__ == "__main__":
    main()
