"""
Generate DGMTS Master Agreement — Code Compliance Inspection Subconsultant Services.
Uses the standard BCC proposal template (header/footer/styles preserved).
Output: Projects/DGMTS/Master Agreement.../
"""
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("pip install python-docx")
    sys.exit(1)

# ── Paths ─────────────────────────────────────────────────────────────────────
TEMPLATE  = Path(r"C:\Users\Kyle Cao\DC Business\Building Code Consulting\DC Code Compliance Proposal Template.docx")
E_SIG     = Path(r"C:\Users\Kyle Cao\DC Business\Building Code Consulting\Logo E-Sig Stamp\E-Sig.jpg")
OUT_DIR   = Path(r"C:\Users\Kyle Cao\DC Business\Building Code Consulting\Projects\DGMTS\Master Agreement – Code Compliance Subconsultant Services")
OUT_DOCX  = OUT_DIR / "DGMTS Master Agreement – Third-Party Code Compliance Inspection Subconsultant Services from BCC.docx"

TODAY       = "February 18, 2026"
CLIENT_FULL = "Dulles Geotechnical and Materials Testing Services, Inc. (DGMTS)"
CLIENT_SHORT = "DGMTS"
PRICE       = 375


# ── Helpers ───────────────────────────────────────────────────────────────────
def _set_black(run):
    run.font.color.rgb = RGBColor(0, 0, 0)

def _replace_para(p, new_text: str) -> None:
    """Replace paragraph text while preserving first-run formatting (font/size/bold/italic)."""
    fmt = {}
    if p.runs:
        r0 = p.runs[0]
        fmt = {
            "bold":      r0.bold,
            "italic":    r0.italic,
            "font_size": r0.font.size,
            "font_name": r0.font.name,
            "underline": r0.underline,
        }
    p.clear()
    run = p.add_run(new_text)
    for attr, val in fmt.items():
        if val is not None:
            if attr == "font_size":   run.font.size = val
            elif attr == "font_name": run.font.name = val
            elif attr == "bold":      run.bold = val
            elif attr == "italic":    run.italic = val
            elif attr == "underline": run.underline = val
    _set_black(run)

def _cell_text(cell, text: str, bold: bool = False, font_size: int = 11) -> None:
    """Set all paragraphs in a cell to the given text (clears existing)."""
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    _set_black(run)

def _add_row(table):
    """Add a new row to a table."""
    new_tr = OxmlElement("w:tr")
    # Copy column structure from first row
    ref_tr = table.rows[0]._tr
    for tc in ref_tr.findall(qn("w:tc")):
        new_tc = OxmlElement("w:tc")
        # Copy cell properties if present
        tcp = tc.find(qn("w:tcPr"))
        if tcp is not None:
            import copy
            new_tc.append(copy.deepcopy(tcp))
        new_p = OxmlElement("w:p")
        new_tc.append(new_p)
        new_tr.append(new_tc)
    table._tbl.append(new_tr)
    return table.rows[-1]

def set_all_text_black(doc):
    black = RGBColor(0, 0, 0)
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.color.rgb = black
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = black
    for section in doc.sections:
        for hf in (section.header, section.footer):
            for p in hf.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = black


# ── Main generator ─────────────────────────────────────────────────────────────
def generate():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(str(TEMPLATE))
    paras = doc.paragraphs

    # ── COVER PAGE ─────────────────────────────────────────────────────────────
    # [3]  Project/type title  (large display text)
    _replace_para(paras[3], "Master Agreement — Code Compliance")

    # [4]  Subtitle
    _replace_para(paras[4], "Inspection Subconsultant Services for DGMTS")

    # [7]  Address street line → client company
    _replace_para(paras[7], CLIENT_FULL)

    # [8]  Address city line → scope descriptor
    _replace_para(paras[8], "Various DC Projects — On-Call Basis")

    # [9]  Attention line
    _replace_para(paras[9], "ATTENTION: Mr. Tariq B. Hamid, P.E. & Mr. Zichang Zhang")

    # [10] Date line
    _replace_para(paras[10], TODAY)

    # [14] "From:" line — keep as-is (Yue Kyle Cao)

    # Use one of the blank paras between attention and "From:" for email info
    # paras[11] is blank — repurpose it for email contact line
    _replace_para(paras[11], "thamid@dullesgeotechnical.com  |  zzhang@dullesgeotechnical.com")

    # ── SCOPE OF WORK intro (para 38) ─────────────────────────────────────────
    _replace_para(
        paras[38],
        'This Master Agreement ("Agreement") sets forth the scope of services, rates, and general terms '
        'under which Building Code Consulting LLC ("BCC") will provide Third-Party Code Compliance '
        'Inspection services as a subconsultant to Dulles Geotechnical and Materials Testing Services, Inc. '
        '("DGMTS") for designated construction projects in the District of Columbia. Once executed, this '
        'Agreement governs all future project assignments without the need for individual proposals per project, '
        'unless the scope materially differs from what is described herein.'
    )

    # ── PROJECT DESCRIPTION — EXHIBIT A (para 73) ─────────────────────────────
    _replace_para(
        paras[73],
        "BCC will serve as the Third-Party Inspection Agency (TPIA) and Subconsultant Inspector for "
        "DGMTS on District of Columbia construction projects that require code compliance inspections as "
        "part of DGMTS's scope of special inspection services. BCC will coordinate directly with the "
        "general contractor and relevant permit holders, while keeping DGMTS informed of all inspection "
        "statuses and outcomes. All inspections will be performed in accordance with the DC Construction "
        "Codes (DCMR 12) and DC Department of Buildings (DOB) regulations. BCC will register each project "
        "in the DOB Tertius online inspection management system as required by DC law, and will upload all "
        "field inspection reports to Tertius after each visit. Copies of reports will be provided to "
        "DGMTS upon request. BCC's role on each project will be to serve as the combo inspection "
        "inspector, assisting DGMTS and the project team with all required code compliance inspections."
    )

    # ── FEE TABLE (Table 0) ────────────────────────────────────────────────────
    fee_tbl = doc.tables[0]

    # Row 0: title spanning all columns (cells are merged in template — only touch cells[0])
    _cell_text(fee_tbl.rows[0].cells[0], "Fee Schedule — Master Rate (All DC Projects)", bold=True)
    # Do NOT call _cell_text on cells[1-3]: they share the same CT_Tc as cells[0] for merged rows

    # Row 1: header row
    hdr = fee_tbl.rows[1].cells
    _cell_text(hdr[0], "Service Description", bold=True)
    _cell_text(hdr[1], "Unit",                bold=True)
    _cell_text(hdr[2], "Rate",                bold=True)
    _cell_text(hdr[3], "",                    bold=False)  # 4th col not used in master rate

    # Rows 2-5: replace existing line-item rows with master rate items
    line_items = [
        ("Combo Inspection Visit\n(Building, Electrical, Mechanical, Plumbing, or Fire Protection)",
         "Per Visit", f"${PRICE}.00", ""),
        ("Re-Inspection / Failed Inspection",
         "Per Visit", f"${PRICE}.00", ""),
        ("Consultation / Administrative Support\n(Beyond standard field reporting)",
         "Per Hour", "$150.00", ""),
        ("",  "", "", ""),   # blank row (was row 5 in original)
    ]
    for i, (svc, unit, rate, _) in enumerate(line_items):
        row = fee_tbl.rows[2 + i]
        _cell_text(row.cells[0], svc)
        _cell_text(row.cells[1], unit)
        _cell_text(row.cells[2], rate)
        _cell_text(row.cells[3], "")

    # Row 6: was the "Total" summary row — repurpose as fee-note row
    last_row = fee_tbl.rows[6]
    _cell_text(last_row.cells[0],
               "Each visit covers up to 3 hours: 1 hr round-trip travel + up to 2 hrs on-site. "
               "Time beyond 2 hrs on-site billed at hourly consultation rate in 30-min increments.")
    _cell_text(last_row.cells[1], "")
    _cell_text(last_row.cells[2], "")
    _cell_text(last_row.cells[3], "")

    # Also update the "Flat rate of $325" line in paragraph 120
    _replace_para(
        paras[120],
        f"Inspection Services: Flat Rate of ${PRICE}.00 per visit (Combo — all applicable disciplines)"
    )

    # ── SIGNATURE TABLE (Table 1) ──────────────────────────────────────────────
    sig_tbl = doc.tables[1]

    # Row 0: party headers
    left  = sig_tbl.rows[0].cells[0]
    right = sig_tbl.rows[0].cells[1]

    # Clear left cell and rebuild for BCC (with e-sig image)
    for p in left.paragraphs:
        p.clear()

    # "For Building Code Consulting LLC,"
    p = left.paragraphs[0]
    run = p.add_run("For Building Code Consulting LLC,")
    run.bold = True
    run.font.size = Pt(11)
    _set_black(run)

    # Blank line then e-sig image
    p2 = left.add_paragraph()
    p2.paragraph_format.space_before = Pt(6)
    if E_SIG.exists():
        run_img = p2.add_run()
        run_img.add_picture(str(E_SIG), height=Inches(0.55))

    # Name / Title / Date
    for line in ["Yue Cao, P.E., MCP", "President", f"Date: {TODAY}"]:
        pn = left.add_paragraph()
        r = pn.add_run(line)
        r.font.size = Pt(11)
        _set_black(r)

    # Clear right cell and rebuild for DGMTS
    for p in right.paragraphs:
        p.clear()

    p = right.paragraphs[0]
    run = p.add_run("Agreed and Accepted by DGMTS:")
    run.bold = True
    run.font.size = Pt(11)
    _set_black(run)

    sig_lines_right = [
        "",
        "Signature:  ___________________________________",
        "Name:  Tariq B. Hamid, P.E.  /  Zichang Zhang",
        "Title:  ___________________________________",
        "Date:  ___________________________________",
    ]
    for line in sig_lines_right:
        pn = right.add_paragraph()
        r = pn.add_run(line)
        r.font.size = Pt(11)
        _set_black(r)

    # Also update para 124 (Signature Block intro text)
    _replace_para(
        paras[124],
        "If you find the terms of this Master Agreement acceptable, please acknowledge your acceptance "
        "by signing below. Once executed, this Agreement will serve as the standing agreement for all "
        "DC code compliance inspection subconsultant work between BCC and DGMTS."
    )

    # ── FINAL: set all text black, save ───────────────────────────────────────
    set_all_text_black(doc)
    doc.save(str(OUT_DOCX))
    print(f"[OK] Saved: {OUT_DOCX}")
    return OUT_DOCX


if __name__ == "__main__":
    out = generate()

    # Convert to PDF
    try:
        from docx2pdf import convert
        pdf_path = str(out).replace(".docx", ".pdf")
        convert(str(out), pdf_path)
        print(f"[OK] PDF:   {pdf_path}")
    except Exception as e:
        print(f"[!]  PDF conversion failed: {e}")
        print("    Open the .docx in Word and File > Save As PDF manually.")
