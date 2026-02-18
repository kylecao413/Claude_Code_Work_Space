"""
Phase 2 Internal Audit: Read generated proposal DOCX; fail if Insomnia, 701 Monroe, red text, or old date.
Returns 0 = pass, 1 = fail.
"""
import re
import sys
from pathlib import Path
from datetime import datetime

def audit(docx_path: Path) -> bool:
    from docx import Document
    from docx.shared import RGBColor
    doc = Document(str(docx_path))
    today = datetime.now().strftime("%m-%d-%Y")
    fail_reasons = []
    def check_red(run):
        try:
            if run.font.color.rgb and run.font.color.rgb != RGBColor(0, 0, 0):
                return True
        except Exception:
            pass
        return False
    for p in doc.paragraphs:
        t = p.text
        if "Insomnia" in t or "Insomnia Cookies" in t:
            fail_reasons.append("Found 'Insomnia' in document")
        if "701 Monroe" in t:
            fail_reasons.append("Found wrong address 701 Monroe")
        if "01-12-2026" in t or (t.strip() and t.strip() != today and re.match(r"01[-\/]12[-\/]2026", t)):
            fail_reasons.append("Found old date (01-12-2026)")
        for run in p.runs:
            if check_red(run):
                fail_reasons.append("Found RED text (placeholder)")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text
                    if "Insomnia" in t or "701 Monroe" in t:
                        fail_reasons.append("Wrong content in table")
                    for run in p.runs:
                        if check_red(run):
                            fail_reasons.append("RED text in table")
    if fail_reasons:
        print("AUDIT FAIL:", "; ".join(fail_reasons), file=sys.stderr)
        return False
    print("AUDIT PASS: No Insomnia, no 701 Monroe, no red text, no old date.")
    return True


if __name__ == "__main__":
    path = Path(r"c:\Users\Kyle Cao\DC Business\Building Code Consulting\Projects\Keller Brothers\St. Joseph's on Capitol Hill – Phase I\St. Joseph's on Capitol Hill – Phase I - Third Party Code Inspection Proposal from BCC.docx")
    if len(sys.argv) > 1:
        path = Path(sys.argv[1])
    sys.exit(0 if audit(path) else 1)
