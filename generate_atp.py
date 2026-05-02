"""
Generate an Authorization to Proceed (ATP) — 1-page per-project document
referencing a parent Master Inspection Contract (MIC).

Architecture:  one signed MIC per client, then one ATP per project.
ATP says "this project, these disciplines, this visit estimate; rates per
Article 4 of MIC-YYYY-NNN."  No new proposal needed, just sign and start.

Numbering: ATP-{YYYY}-{NNN} sequential globally across all clients.
Output:    Master Contract And ATP/<Client>/ATP/ATP-YYYY-NNN <Project>.docx
"""
import argparse
import re
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("pip install python-docx")
    sys.exit(1)

ROOT  = Path(r"C:\Users\Kyle Cao\DC Business\Building Code Consulting\Master Contract And ATP")
E_SIG = Path(r"C:\Users\Kyle Cao\DC Business\Building Code Consulting\Logo E-Sig Stamp\E-Sig.jpg")

ALL_DISCIPLINES = ["Building", "Mechanical", "Electrical", "Plumbing", "Fire Protection"]


def find_parent_mic(client: str) -> tuple[str, Path]:
    """Locate the client's MIC docx and return (mic_no, path). Errors if missing."""
    client_dir = ROOT / client
    if not client_dir.exists():
        print(f"[ERR] No client folder at {client_dir}")
        print(f"      Generate the MIC first: python generate_mic_master_contract.py --client \"{client}\" ...")
        sys.exit(1)
    for f in client_dir.glob("*.docx"):
        if "Master Proposal" in f.name or "MIC" in f.name:
            d = Document(str(f))
            for p in d.paragraphs[:15]:
                m = re.search(r"MIC-\d{4}-\d{3}", p.text)
                if m:
                    return m.group(0), f
    print(f"[ERR] No MIC found in {client_dir}")
    sys.exit(1)


def next_atp_number() -> str:
    """Scan all ATP/ subfolders across clients for highest ATP-YYYY-NNN; return next."""
    used = set()
    for f in ROOT.rglob("ATP/*.docx"):
        m = re.search(r"ATP-(\d{4})-(\d{3})", f.name)
        if m:
            used.add(int(m.group(2)))
    n = (max(used) + 1) if used else 1
    return f"ATP-{datetime.today().year}-{n:03d}"


def _set_black(run):
    run.font.color.rgb = RGBColor(0, 0, 0)


def _add(doc, text, bold=False, size=11, align=None, space_after=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    _set_black(r)
    return p


def _kv_row(table, label, value):
    """Add a label/value row to a 2-col table."""
    row = table.add_row()
    for cell, txt, bold in [(row.cells[0], label, True), (row.cells[1], value, False)]:
        for p in cell.paragraphs:
            p.clear()
        p = cell.paragraphs[0]
        r = p.add_run(txt)
        r.bold = bold
        r.font.size = Pt(11)
        _set_black(r)


def generate(client: str, project_name: str, project_address: str, permit_no: str,
             description: str, disciplines: list[str], est_visits: int,
             single_rate: int = 325, combo_rate: int | None = None,
             timeline: str | None = None, atp_no: str | None = None) -> Path:

    mic_no, mic_path = find_parent_mic(client)
    atp_no = atp_no or next_atp_number()
    today_long = datetime.today().strftime("%B %d, %Y")
    today_short = datetime.today().strftime("%#m/%#d/%Y") if sys.platform == "win32" else datetime.today().strftime("%-m/%-d/%Y")

    # Visit type for pricing math
    is_combo = len(disciplines) >= 2
    rate_per_visit = (combo_rate if is_combo and combo_rate is not None else single_rate)
    visit_type = "combination" if is_combo and combo_rate is not None else "single-trade"
    est_total = est_visits * rate_per_visit

    out_dir = ROOT / client / "ATP"
    out_dir.mkdir(parents=True, exist_ok=True)
    safe_proj = re.sub(r"[\\/:*?\"<>|]", "_", project_name)[:80]
    out_file = out_dir / f"{atp_no} {safe_proj}.docx"

    d = Document()

    # Header
    _add(d, "AUTHORIZATION TO PROCEED (ATP)", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add(d, "Issued under a BCC Master Inspection Contract", bold=False, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))

    # ATP / MIC reference table
    ref_tbl = d.add_table(rows=0, cols=2)
    ref_tbl.autofit = True
    _kv_row(ref_tbl, "ATP Reference No.:", atp_no)
    _kv_row(ref_tbl, "Parent MIC:", mic_no)
    _kv_row(ref_tbl, "Issue Date:", today_long)
    _kv_row(ref_tbl, "Client:", client)
    _add(d, "")

    # Project details
    _add(d, "PROJECT DETAILS", bold=True, size=12)
    proj_tbl = d.add_table(rows=0, cols=2)
    _kv_row(proj_tbl, "Project Name:", project_name)
    _kv_row(proj_tbl, "Project Address:", project_address)
    _kv_row(proj_tbl, "Permit Number:", permit_no)
    if timeline:
        _kv_row(proj_tbl, "Timeline:", timeline)
    _kv_row(proj_tbl, "Project Description:", description)
    _add(d, "")

    # Scope of inspections
    _add(d, "SCOPE OF INSPECTIONS", bold=True, size=12)
    for disc in ALL_DISCIPLINES:
        mark = "[X]" if disc in disciplines else "[  ]"
        _add(d, f"   {mark}  {disc}", size=11, space_after=Pt(2))
    _add(d, "")

    # Pricing
    _add(d, "PRICING (per Article 4 of " + mic_no + ")", bold=True, size=12)
    _add(d, f"Visit Type: {visit_type.title()} ({len(disciplines)} discipline{'s' if len(disciplines)!=1 else ''} per visit)")
    _add(d, f"Rate: ${rate_per_visit}.00 per visit")
    _add(d, f"Estimated Visits: {est_visits}")
    _add(d, f"Estimated Total: ${est_total:,}.00", bold=True)
    _add(d, "Billing is based on actual visits completed — flat rate per visit actually performed, "
            "never billed based on upfront estimate. Visits beyond the estimate are billed at the same per-visit rate.",
         size=10)
    _add(d, "")

    # Authorization & signatures
    _add(d, "AUTHORIZATION", bold=True, size=12)
    _add(d, f"BCC is hereby authorized to commence Third-Party Code Compliance Inspection services on the above "
            f"project under the terms of {mic_no}. All general terms, conditions, and process requirements of "
            f"{mic_no} apply to this ATP. Inspection requests will be coordinated directly with the Client's site "
            f"superintendent or project manager.")
    _add(d, "")

    # Signature table
    sig_tbl = d.add_table(rows=1, cols=2)
    left = sig_tbl.rows[0].cells[0]
    right = sig_tbl.rows[0].cells[1]

    # Left: BCC with e-sig
    for p in left.paragraphs:
        p.clear()
    pl = left.paragraphs[0]
    rl = pl.add_run("For Building Code Consulting LLC,")
    rl.bold = True; rl.font.size = Pt(11); _set_black(rl)
    p2 = left.add_paragraph()
    p2.paragraph_format.space_before = Pt(6)
    if E_SIG.exists():
        p2.add_run().add_picture(str(E_SIG), height=Inches(0.55))
    for line in ["Yue Cao, P.E., MCP", "President", f"Date: {today_short}"]:
        pn = left.add_paragraph()
        rn = pn.add_run(line)
        rn.font.size = Pt(11); _set_black(rn)

    # Right: Client acknowledgment
    for p in right.paragraphs:
        p.clear()
    pr = right.paragraphs[0]
    rr = pr.add_run(f"Acknowledged by {client}:")
    rr.bold = True; rr.font.size = Pt(11); _set_black(rr)
    for line in ["", "Signature: ___________________________",
                 "Name: ___________________________",
                 "Title: ___________________________",
                 "Date: ___________________________"]:
        pn = right.add_paragraph()
        rn = pn.add_run(line)
        rn.font.size = Pt(11); _set_black(rn)

    d.save(str(out_file))
    print(f"[OK] {atp_no}  →  {out_file}")
    print(f"     parent MIC: {mic_no}  ({mic_path.name})")
    return out_file


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--client", required=True, help='Client name matching their folder under Master Contract And ATP/')
    ap.add_argument("--project", required=True, help="Project name")
    ap.add_argument("--address", required=True)
    ap.add_argument("--permit", required=True)
    ap.add_argument("--description", required=True)
    ap.add_argument("--disciplines", required=True,
                    help='Comma-separated subset of: Building,Mechanical,Electrical,Plumbing,Fire Protection')
    ap.add_argument("--visits", type=int, required=True, help="Estimated visit count")
    ap.add_argument("--single-rate", type=int, default=325)
    ap.add_argument("--combo-rate", type=int, default=None,
                    help="If client's MIC has a two-tier rate, set this to the combo rate")
    ap.add_argument("--timeline", default=None)
    ap.add_argument("--atp-no", default=None)
    args = ap.parse_args()

    discs = [d.strip() for d in args.disciplines.split(",") if d.strip()]
    bad = [d for d in discs if d not in ALL_DISCIPLINES]
    if bad:
        print(f"[ERR] Unknown discipline(s): {bad}.  Valid: {ALL_DISCIPLINES}")
        sys.exit(1)

    generate(args.client, args.project, args.address, args.permit, args.description,
             discs, args.visits, args.single_rate, args.combo_rate,
             args.timeline, args.atp_no)
