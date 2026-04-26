"""
基于 Word 模板生成 DC 第三方检测服务提案，并按定价逻辑计算 Combo Inspection 单价。
输出到 ../Projects/[Client Name]/[Project Name]/；可选在终端打印定价摘要表供 Yue 确认后再生成。
"""
import os
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

try:
    from dotenv import load_dotenv
    load_dotenv(Path(__file__).resolve().parent / ".env")
except ImportError:
    pass

BASE_DIR = Path(__file__).resolve().parent
# 项目数据区；模板优先使用 Building Code Consulting 根目录下你提供的模板
PROJECTS_ROOT = BASE_DIR.parent / "Projects"
TEMPLATE_DIR = BASE_DIR / "BuildingConnected" / "templates"
# 你提供的模板路径（Building Code Consulting 根目录）
TEMPLATE_CODE_COMPLIANCE = Path(r"c:\Users\Kyle Cao\DC Business\Building Code Consulting\DC Code Compliance Proposal Template.docx")
TEMPLATE_PLAN_REVIEW = Path(r"c:\Users\Kyle Cao\DC Business\Building Code Consulting\DC Plan Review Proposal Template.docx")
# 回退：本地 templates 目录
DEFAULT_TEMPLATE = TEMPLATE_DIR / "DC Code Compliance Proposal Template.docx"


def get_template_path(template_type: str = "code_compliance") -> Path:
    """优先使用你提供的模板路径，否则回退到 BuildingConnected/templates/。"""
    if template_type == "plan_review":
        if TEMPLATE_PLAN_REVIEW.exists():
            return TEMPLATE_PLAN_REVIEW
        return TEMPLATE_DIR / "DC Plan Review Proposal Template.docx"
    if TEMPLATE_CODE_COMPLIANCE.exists():
        return TEMPLATE_CODE_COMPLIANCE
    return DEFAULT_TEMPLATE

# 定价逻辑（$/visit Combo Inspection）
PRICING_TIERS = {
    "key_large": 295,      # 极大型项目/大客户
    "regular": (300, 350), # 常规/中型回头客
    "small_repeat": (350, 375),  # 小型重复客户
    "one_time": (375, 400),      # 罕见/一次性客户
}


def sanitize_dirname(s: str) -> str:
    """用于生成 Client/Project 文件夹名的安全字符串。"""
    s = re.sub(r'[<>:"/\\|?*]', "_", (s or "").strip())
    return s[:80].strip(" .") or "Unknown"


def suggest_tier(project: dict) -> tuple[str, int, str]:
    """
    根据项目信息建议定价档位与单价，及理由。
    project 可含: client, name, size_sqft, is_repeat, is_key_account 等。
    """
    client = (project.get("client") or "").strip()
    name = (project.get("name") or "").strip()
    size = project.get("size_sqft") or project.get("scope_notes") or ""
    is_repeat = project.get("is_repeat", False)
    is_key = project.get("is_key_account", False)
    # 极大型/大客户
    if is_key or (size and "large" in str(size).lower()) or ("marriott" in name.lower() or "flagship" in name.lower()):
        return "key_large", PRICING_TIERS["key_large"], "极大型/大客户 $295/visit"
    # 常规/中型回头客
    if is_repeat and not client.startswith("Small"):
        return "regular", (PRICING_TIERS["regular"][0] + PRICING_TIERS["regular"][1]) // 2, "常规/中型回头客 $300–350/visit"
    # 小型重复
    if is_repeat:
        return "small_repeat", (PRICING_TIERS["small_repeat"][0] + PRICING_TIERS["small_repeat"][1]) // 2, "小型重复客户 $350–375/visit"
    # 一次性
    return "one_time", (PRICING_TIERS["one_time"][0] + PRICING_TIERS["one_time"][1]) // 2, "一次性客户 $375–400/visit"


def print_pricing_summary(project: dict, tier: str, price_per_visit: int, reason: str, est_visits: int = 12):
    """在终端打印定价摘要表，供 Yue 手机回复「同意」或修改。"""
    total = price_per_visit * est_visits
    print("\n" + "=" * 60)
    print("定价摘要（请确认或回复修改）")
    print("=" * 60)
    print(f"  项目: {project.get('name', 'N/A')}")
    print(f"  客户: {project.get('client', 'N/A')}")
    print(f"  档位: {tier} | 理由: {reason}")
    print(f"  建议单价: ${price_per_visit}/visit (Combo Inspection)")
    print(f"  预估次数: {est_visits} visits → 合计约 ${total}")
    print("=" * 60)
    print("回复「同意」将按此生成提案；或指定单价如 320 再生成。\n")


def get_proposal_output_dir(client_name: str, project_name: str) -> Path:
    """../Projects/[Client]/[Project]/，不存在则创建。"""
    client = sanitize_dirname(client_name)
    project = sanitize_dirname(project_name)
    out_dir = PROJECTS_ROOT / client / project
    out_dir.mkdir(parents=True, exist_ok=True)
    return out_dir


def _split_address(address: str) -> tuple[str, str]:
    """Split 'Street, City, State ZIP' into (street_line, city_state_line).
    e.g. '20 F Street NW, Suite 550, Washington, DC 20001'
         → ('20 F Street NW, Suite 550', 'Washington, DC 20001')
    Uses 'Washington, DC' as the primary split marker for DC addresses,
    then falls back to rsplit for other cities.
    """
    addr = (address or "").strip()
    if not addr:
        return "", "Washington, DC"

    # Primary: look for "Washington, DC" as the city marker (more specific than bare "Washington")
    dc_idx = addr.find("Washington, DC")
    if dc_idx > 0:
        street = addr[:dc_idx].rstrip(", ").strip()
        city = addr[dc_idx:].strip()
        if street and city:
            return street, city

    # Fallback: rsplit for "Street, [Suite], City, State ZIP" patterns
    parts = addr.rsplit(",", 2)
    if len(parts) == 3:
        return parts[0].strip(), ", ".join(parts[1:]).strip()
    if len(parts) == 2:
        return parts[0].strip(), parts[1].strip()
    return addr, "Washington, DC"


def _generate_default_description(project: dict) -> str:
    """Generate a formal project description when none is provided."""
    name = project.get("name", "the project")
    address = project.get("address", "the project location")
    return (
        f"Third-party code compliance inspection services for {name} "
        f"at {address}. Building Code Consulting LLC will perform all applicable "
        "building, mechanical, electrical, plumbing, and fire protection "
        "inspections in accordance with the DC Building Code and all applicable "
        "DC Department of Buildings regulations."
    )


def _clean_description(desc_text: str, address: str, client: str) -> str:
    """Strip bid deadline sentences + any caller-supplied 'BCC's role' / address repeats,
    then append the canonical BCC closing sentence exactly once."""
    import re
    # Split into sentences
    sentences = re.split(r'(?<=[.!?])\s+', desc_text.strip())
    dropped_patterns = [
        re.compile(r'bid\s*(due|deadline|date)|due\s*date|deadline', re.IGNORECASE),
        # Any sentence restating BCC's role — caller sometimes writes their own version.
        # Drop them so we emit only the canonical closing below.
        re.compile(r"BCC[’']?s?\s+role", re.IGNORECASE),
        re.compile(r"BCC\s+will\s+serve\s+as", re.IGNORECASE),
        # Any sentence that re-states the project address — avoid duplicating the closing address line.
        re.compile(r"(the\s+)?project\s+address\s+is", re.IGNORECASE),
    ]
    filtered = [
        s for s in sentences
        if not any(p.search(s) for p in dropped_patterns)
    ]
    cleaned = ' '.join(filtered).strip().rstrip('.')

    # Don't append ", Washington DC" if address already contains it
    if re.search(r'Washington[,\s]*DC', address, re.IGNORECASE):
        addr_part = address
    else:
        addr_part = f"{address}, Washington DC"

    # Standard BCC closing sentence
    closing = (
        f"The project address is {addr_part}. "
        f"BCC's role on the project will be to serve as the combo inspection inspector, "
        f"assisting {client} with all required inspections."
    )

    # Append closing once (we've stripped any caller version above)
    if "combo inspection inspector" not in cleaned:
        cleaned = (cleaned + ". " if cleaned else "") + closing

    return cleaned


def _replace_para_preserving_format(p, new_text: str) -> None:
    """Replace all text in a paragraph with new_text, preserving the first run's font formatting.
    Avoids the font-size-collapse bug caused by p.clear() + bare p.add_run()."""
    fmt = {}
    if p.runs:
        r0 = p.runs[0]
        fmt['bold'] = r0.bold
        fmt['italic'] = r0.italic
        fmt['font_size'] = r0.font.size
        fmt['font_name'] = r0.font.name
        fmt['underline'] = r0.underline
    p.clear()
    run = p.add_run(new_text)
    if fmt.get('bold') is not None:
        run.bold = fmt['bold']
    if fmt.get('italic') is not None:
        run.italic = fmt['italic']
    if fmt.get('font_size') is not None:
        run.font.size = fmt['font_size']
    if fmt.get('font_name') is not None:
        run.font.name = fmt['font_name']
    if fmt.get('underline') is not None:
        run.underline = fmt['underline']


def _collapse_empty_paragraphs(doc, max_consecutive: int = 2) -> None:
    """Remove excessive consecutive empty Normal paragraphs (keep ≤ max_consecutive)."""
    to_remove = []
    empty_count = 0
    for p in doc.paragraphs:
        if not p.text.strip() and p.style.name in ("Normal", ""):
            empty_count += 1
            if empty_count > max_consecutive:
                to_remove.append(p)
        else:
            empty_count = 0
    WP_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    for p in to_remove:
        # Never remove paragraphs containing section breaks — they control
        # footer linkage, page numbering, and section properties.
        pPr = p._element.find(f'{WP_NS}pPr')
        if pPr is not None and pPr.find(f'{WP_NS}sectPr') is not None:
            continue
        if p._element.find(f'{WP_NS}sectPr') is not None:
            continue
        try:
            p._element.getparent().remove(p._element)
        except (AttributeError, ValueError):
            pass


def fill_template(
    template_path: Path,
    out_path: Path,
    project: dict,
    price_per_visit: int,
    est_visits: int = 12,
    scope_notes: str = "",
) -> bool:
    """
    使用 python-docx 打开模板，替换占位符（如 {{Client}}、{{Project}}、{{PricePerVisit}} 等），保存到 out_path。
    """
    try:
        from docx import Document
    except ImportError:
        print("请安装 python-docx: pip install python-docx")
        return False
    if template_path.exists():
        doc = Document(str(template_path))
    else:
        print(f"未找到模板 {template_path}，使用内置简易模板。请将正式模板放入 BuildingConnected/templates/ 以获得正确版式。")
        doc = Document()
        doc.add_paragraph("Third-Party Code Compliance Inspection Proposal")
        doc.add_paragraph("Client: {{Client}}")
        doc.add_paragraph("Project: {{Project}}")
        doc.add_paragraph("Price per visit (Combo): ${{PricePerVisit}}")
        doc.add_paragraph("Estimated visits: {{EstVisits}} | Total: ${{Total}}")
        doc.add_paragraph("Scope: {{ScopeNotes}}")
        doc.add_paragraph("Building Code Consulting LLC – DC Third-Party agency. Yue Cao (PE, MCP).")
    total = price_per_visit * est_visits
    today = datetime.now().strftime("%m-%d-%Y")

    # Split address into street line and city/state line for cover page
    full_addr = project.get("address", "")
    addr_street, addr_city = _split_address(full_addr)

    replacements = {
        # Existing placeholders
        "{{Client}}": project.get("client", ""),
        "{{Project}}": project.get("name", ""),
        "{{ProjectName}}": project.get("name", ""),
        "{{Attention}}": project.get("attention", project.get("contact_name", "")),
        "{{Address}}": full_addr,
        "{{AddressStreet}}": addr_street,
        "{{AddressCity}}": addr_city,
        "{{ProjectAddress}}": full_addr,
        "{{ContactName}}": project.get("attention", project.get("contact_name", "")),
        "{{ContactEmail}}": project.get("contact_email", ""),
        "{{PricePerVisit}}": str(price_per_visit),
        "{{EstVisits}}": str(est_visits),
        "{{Total}}": str(total),
        "{{ScopeNotes}}": scope_notes or project.get("description", "")[:2000],
        "{{BuildingCodeConsulting}}": "Building Code Consulting LLC",
        "{{DC Third-Party}}": "DC Third-Party agency",
        "{{Yue Cao}}": "Yue Cao (PE, MCP)",
        "{{Date}}": today,
        "{{DateLong}}": datetime.now().strftime("%B %d, %Y"),
        # Template leftovers (hardcoded old-project values in the .docx)
        "Insomnia Cookies Renovation": project.get("name", ""),
        "Insomnia Cookies": project.get("name", ""),
        "Cox & Company, LLC": project.get("client", ""),
        "Cox & Company": project.get("client", ""),
        "Bryan Kerr": project.get("attention", project.get("contact_name", "")),
        # Cover page street line → street only (NOT full address, city line handled separately)
        "701 Monroe Street NE": addr_street,
        "701 Monroe ST NE": addr_street,
        "701 Monroe": addr_street,
        "01-12-2026": today,
        "01/12/2026": today,
    }

    def replace_in_paragraph(para):
        for run in para.runs:
            for k, v in replacements.items():
                if k in run.text:
                    run.text = run.text.replace(k, str(v))

    # First pass: run-level replacement
    for p in doc.paragraphs:
        replace_in_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)

    # Second pass: paragraph-level replace to catch text split across runs
    # Uses _replace_para_preserving_format to avoid losing run-level font size overrides
    for p in doc.paragraphs:
        full = p.text
        changed = False
        for find_str, repl in replacements.items():
            if find_str in full:
                full = full.replace(find_str, str(repl))
                changed = True
        if changed:
            _replace_para_preserving_format(p, full)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full = p.text
                    changed = False
                    for find_str, repl in replacements.items():
                        if find_str in full:
                            full = full.replace(find_str, str(repl))
                            changed = True
                    if changed:
                        _replace_para_preserving_format(p, full)

    # Third pass: fix paragraphs with stubborn split-run values
    # Determine best description text once; strip bid deadline info and append standard closing
    raw_desc = scope_notes or project.get("description", "") or _generate_default_description(project)
    desc_text = _clean_description(
        raw_desc,
        address=full_addr or project.get("address", "the project location"),
        client=project.get("client", "the client"),
    )
    for p in doc.paragraphs:
        full = p.text
        if "tenant fit out" in full and "AHUs" in full:
            # Old Insomnia Cookies Exhibit A description — replace with this project's description
            _replace_para_preserving_format(p, desc_text)
        elif full.strip() == "Washington DC":
            # City/state line below street address on cover page — use city portion only
            _replace_para_preserving_format(p, addr_city or "Washington, DC")
        elif "01-12-2026" in full or full.strip() == "01/12/2026":
            _replace_para_preserving_format(p, today)
        elif "Flat rate of $325" in full or "(Flat rate of $325" in full:
            _replace_para_preserving_format(
                p,
                f"Inspection Services Estimated: (Flat rate of ${price_per_visit}/visit)  "
                f"Note: BCC's invoice is fully based on the number of actual combo-inspection visits "
                f"conducted. The visit count and fee total shown below are an estimate for reference only; "
                f"each visit performed is billed separately at the flat per-visit rate, and the estimate "
                f"is neither a cap nor a commitment to a specific visit count."
            )

    # Fourth pass: ensure Project Description heading is followed by description text
    # (handles case where Exhibit A paragraph was already empty in template after replacements)
    paragraphs = doc.paragraphs
    for i, p in enumerate(paragraphs):
        if p.text.strip() == "Project Description - Exhibit A" and "Heading" in p.style.name:
            # Check if the third pass already added content within the next 5 paragraphs
            has_desc = any(
                paragraphs[j].text.strip()
                for j in range(i + 1, min(i + 6, len(paragraphs)))
            )
            if not has_desc:
                # No content found — fill the next paragraph with the description
                if i + 1 < len(paragraphs):
                    _replace_para_preserving_format(paragraphs[i + 1], desc_text)
            break  # Only one "Project Description - Exhibit A" heading expected

    # Fifth pass: set keep_with_next on non-empty headings to prevent orphan section titles;
    # force page break before major section headings (Index, Scope of Work, Exhibit*, all H1
    # after the first); convert empty Heading paragraphs to Normal to avoid spacing artifacts.
    _PAGE_BREAK_KEYWORDS = ("exhibit", "index", "table of contents", "scope of work",
                            "terms and conditions", "fee schedule", "schedule of values")
    first_h1_seen = False
    for p in doc.paragraphs:
        if "Heading" in p.style.name:
            if p.text.strip():
                p.paragraph_format.keep_with_next = True
                txt_lower = p.text.strip().lower()
                is_h1 = p.style.name in ("Heading 1", "Heading1")
                # Page break before every Heading 1 except the very first (cover page title),
                # and before any heading whose text matches a known section keyword.
                needs_break = any(kw in txt_lower for kw in _PAGE_BREAK_KEYWORDS)
                if is_h1:
                    if first_h1_seen:
                        needs_break = True
                    else:
                        first_h1_seen = True
                if needs_break:
                    p.paragraph_format.page_break_before = True
            else:
                # Empty heading — demote to Normal so it doesn't add extra heading-style space
                try:
                    p.style = doc.styles["Normal"]
                except Exception:
                    pass

    # Sixth pass: keep signature block together on one page.
    # Find consecutive paragraphs that form the signature section and mark keep_with_next.
    import re as _re
    _SIG_MARKERS = frozenset([
        "authorized signature", "signature:", "accepted by", "agreed:", "sign here",
    ])
    _FIELD_MARKERS = frozenset(["name:", "title:", "company:", "date:", "phone:", "email:"])
    _ALL_SIG = _SIG_MARKERS | _FIELD_MARKERS
    paras = list(doc.paragraphs)
    i = 0
    while i < len(paras):
        txt_l = paras[i].text.strip().lower()
        if any(m in txt_l for m in _ALL_SIG):
            # Walk back one paragraph to include any intro line (e.g. "For BCC:")
            block_start = max(0, i - 1)
            j = i
            while j < len(paras):
                tl = paras[j].text.strip().lower()
                is_sig_line = (
                    any(m in tl for m in _ALL_SIG)
                    or bool(_re.search(r'_{3,}', paras[j].text))
                    or not paras[j].text.strip()
                )
                if is_sig_line and j - block_start < 20:
                    j += 1
                else:
                    break
            # Mark all paragraphs in block except the last as keep_with_next
            for k in range(block_start, j - 1):
                paras[k].paragraph_format.keep_with_next = True
            i = j
        else:
            i += 1

    # Seventh pass: fix the WHOLE fee table — per-row Price/Visits/Total AND the summary row
    # Template structure:
    #   row 0: "Code Inspections Estimated Fee" (header, merged)
    #   row 1: "Services | Visits | Price Per | Total" (column headers)
    #   row 2: Underground plumbing + water column test
    #   row 3: MEP Combo Close-in (rough-in + framing)
    #   row 4: Air seal close-in / Insulation
    #   row 5: Combo Final Inspections
    #   row 6 (last): blank | <visits> | "Total" | "$<total>"
    def _allocate_visits(total: int, n_data_rows: int) -> list[int]:
        """把 total visits 分到 data rows: [UG plumbing, rough-in, close-in, final].
        Aggressive 哲学 (Kyle 2026-04-22):
          - final 默认 1 次；extras 优先加到 rough-in + close-in（实际施工按层/墙/屋顶分批的是这两阶段）
          - rough-in + close-in 都到 4 次后才给 final 加（超大 multi-phase 项目）
        """
        alloc = [0] * n_data_rows
        if total <= 0 or n_data_rows == 0:
            return alloc
        if n_data_rows == 4:
            if total == 1:
                alloc[3] = 1
            elif total == 2:
                alloc[1] = alloc[3] = 1
            elif total == 3:
                alloc[1] = alloc[2] = alloc[3] = 1
            elif total == 4:
                alloc = [1, 1, 1, 1]
            else:  # > 4: extras 轮流给 rough-in (idx 1) 和 close-in (idx 2)
                alloc = [1, 1, 1, 1]
                extras = total - 4
                # 轮流添加到 rough-in + close-in（alternating），两个都到 4 后再给 final
                cursor = 1  # start with rough-in
                while extras > 0:
                    if alloc[1] < 4 or alloc[2] < 4:
                        # pick the one that's smaller (alternating effect)
                        if alloc[1] <= alloc[2] and alloc[1] < 4:
                            alloc[1] += 1
                        elif alloc[2] < 4:
                            alloc[2] += 1
                        elif alloc[1] < 4:
                            alloc[1] += 1
                        extras -= 1
                    else:
                        alloc[3] += extras
                        extras = 0
        else:
            remaining = total
            for i in range(n_data_rows - 1):
                if remaining > 0:
                    alloc[i] = 1
                    remaining -= 1
            alloc[-1] = max(1, remaining)
        return alloc

    for table in doc.tables:
        rows = table.rows
        if len(rows) < 3:
            continue
        last_row = rows[-1]
        last_cells = last_row.cells
        if len(last_cells) < 3 or "Total" not in last_cells[2].text:
            continue  # 不是 fee table

        # Data rows = row 2 .. last-1 (skip header + column-header rows, skip total row)
        # 但模板前两行可能是 1 header + 1 column-header, 保险起见找第一个含 "$" 的 data row 开始
        first_data_idx = None
        for ri in range(2, len(rows) - 1):
            row_text = " ".join(c.text for c in rows[ri].cells)
            if "$" in row_text:
                first_data_idx = ri
                break
        if first_data_idx is None:
            continue
        data_rows = rows[first_data_idx:-1]
        n_data = len(data_rows)
        alloc = _allocate_visits(est_visits, n_data)

        # Expand fee-table service descriptions with conditional sub-items.
        # Kyle 2026-04-22, option B:
        #   Row 3 (MEP Close-in): + sprinkler hydro/flush test (if required by code)
        #   Row 4 (Air seal / Insulation): + exterior wall sheathing (if required by code)
        #   Row 5 (Combo Final): + fire alarm + fire suppression acceptance test (if required by code)
        # Use "if required by code" so GCs on projects without those systems aren't scared.
        CLOSE_IN_SUFFIX = ", including sprinkler hydro/flush test (if required by code)"
        AIR_SEAL_SUFFIX = ", including exterior wall sheathing (if required by code)"
        FINAL_SUFFIX = ", including fire alarm and fire suppression system acceptance testing (if required by code)"
        for dr in data_rows:
            if len(dr.cells) < 1:
                continue
            svc_cell = dr.cells[0]
            base_txt = (svc_cell.text or "").strip()
            if "including" in base_txt:
                continue  # already expanded
            if "MEP rough-in and framing" in base_txt and "Close-in" in base_txt:
                suffix = CLOSE_IN_SUFFIX
            elif "Air seal" in base_txt or "Insulation" in base_txt:
                suffix = AIR_SEAL_SUFFIX
            elif base_txt.startswith("Combo Final Inspection"):
                suffix = FINAL_SUFFIX
            else:
                continue
            for para in svc_cell.paragraphs:
                if para.text.strip() and para.text.strip() in base_txt:
                    _replace_para_preserving_format(para, base_txt + suffix)
                    break

        # Update each data row
        for i, dr in enumerate(data_rows):
            cells = dr.cells
            if len(cells) < 4:
                continue
            v = alloc[i]
            # Visits column (cell[1])
            for para in cells[1].paragraphs:
                txt = para.text.strip()
                if txt.isdigit() or not txt:
                    _replace_para_preserving_format(para, str(v) if v > 0 else "—")
                    break
            # Price Per column (cell[2])
            for para in cells[2].paragraphs:
                if "$" in para.text:
                    _replace_para_preserving_format(para, f"${price_per_visit}" if v > 0 else "—")
                    break
            # Total column (cell[3])
            for para in cells[3].paragraphs:
                if "$" in para.text:
                    row_total = v * price_per_visit
                    _replace_para_preserving_format(para, f"${row_total:,}" if v > 0 else "—")
                    break

        # Total row (last)
        for para in last_cells[1].paragraphs:
            if para.text.strip().lstrip("$").isdigit():
                _replace_para_preserving_format(para, str(est_visits))
                break
        if len(last_cells) > 3:
            for para in last_cells[3].paragraphs:
                if para.text.strip().startswith("$"):
                    _replace_para_preserving_format(para, f"${price_per_visit * est_visits:,}")
                    break

    # Eighth pass: discipline applicability (mark rows (not applicable) if outside scope)
    disciplines = (project.get("disciplines") or set())
    if disciplines and disciplines != {"building", "mechanical", "electrical", "plumbing", "fire_protection"}:
        import re
        DISCIPLINE_MAP = [
            ("building", "Building"),
            ("mechanical", "Mechanical"),
            ("electrical", "Electrical"),
            ("plumbing", "Plumbing"),
            ("fire_protection", "Fire Protection"),
        ]
        paras = list(doc.paragraphs)
        for i, p in enumerate(paras):
            txt_stripped = p.text.strip()
            for key, label in DISCIPLINE_MAP:
                if re.match(rf"^{re.escape(label)}\s*\(\s*applicable\s*\)", txt_stripped, re.IGNORECASE):
                    if key not in disciplines:
                        new_heading = re.sub(r"\(\s*applicable\s*\)", "(not applicable)", p.text, flags=re.IGNORECASE)
                        _replace_para_preserving_format(p, new_heading)
                        # Replace the description paragraph (next non-empty) with N/A note
                        for j in range(i + 1, min(i + 4, len(paras))):
                            if paras[j].text.strip():
                                _replace_para_preserving_format(
                                    paras[j],
                                    "Not within project scope — this discipline does not apply to this project."
                                )
                                break
                    break

    # Eighth.5 pass: Fee Schedule — compute hourly rates that back out the flat per-visit fee
    # Logic (Kyle 2026-04-20): flat fee = 2 hrs inspector + 1 PIC review per visit
    #   price >= 350: inspector $150/hr, PIC = price - $300
    #   $295-$325:    inspector $125/hr, PIC = price - $250
    #   < $295:       inspector $100/hr, PIC = price - $200
    if price_per_visit >= 350:
        _ins_rate = 150
    elif price_per_visit >= 295:
        _ins_rate = 125
    else:
        _ins_rate = 100
    _pic_charge = price_per_visit - 2 * _ins_rate
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("Hourly Rates:") and len(t) < 120:
            _replace_para_preserving_format(
                p,
                f"Hourly Rates (for additional on-site time beyond the 2 hours included in each "
                f"per-visit flat fee; minimum billing period is 30 minutes):  "
                f"Code Inspector — ${_ins_rate}/hour.  "
                f"Professional In Charge (PIC) review — ${_pic_charge} per visit (included in flat fee).  "
                f"(The ${price_per_visit}/visit flat fee covers 2 hours of inspector on-site time at "
                f"${_ins_rate}/hr = ${_ins_rate*2} plus the ${_pic_charge} PIC review charge. "
                f"Evenings and weekends are billed at the same rate — no surcharge.)"
            )
            break

    # Eighth.6 pass: Inspection Visits — reconcile "flat rate" with old "3 hours" language
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("Each inspection visit entitles the Client to 3 hours"):
            _replace_para_preserving_format(
                p,
                f"Each inspection visit is billed at the flat per-visit rate in Exhibit C (${price_per_visit}/visit). "
                f"The visit fee covers up to 2 hours of on-site inspection time plus travel. "
                f"Additional on-site time beyond 2 hours is billed per the Fee Schedule in 30-minute increments."
            )
            break
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("Any time beyond 2 hours of onsite inspection"):
            # Already covered in the combined paragraph above — blank this line
            _replace_para_preserving_format(p, "")
            break

    # (Removed Eighth.7 Optional Scopes expansion per Kyle 2026-04-20:
    #  - Plan Review on inspection proposal = conflict of interest
    #  - Fire Alarm Testing is already part of Fire Protection
    #  - BCC doesn't provide Special Inspections — don't mention
    #  Template's existing "Elevator" optional stays as-is.)

    # Ninth pass: signature block — add actual signature / name / title / date lines
    for table in doc.tables:
        if len(table.rows) != 1 or len(table.rows[0].cells) != 2:
            continue
        c0 = table.rows[0].cells[0]
        c1 = table.rows[0].cells[1]
        t0 = (c0.text or "").strip()
        t1 = (c1.text or "").strip()
        if not ("Building Code Consulting" in t0 and "For Client" in t1):
            continue
        # BCC side: add sig lines after existing "Yue Cao – President" text
        bcc_lines = [
            "",
            "Signature: _________________________________",
            "Printed Name: Yue (Kyle) Cao, PE, MCP",
            "Title: President",
            "Date: _________________________________",
        ]
        client_lines = [
            "",
            "Signature: _________________________________",
            "Printed Name: _________________________________",
            "Title: _________________________________",
            "Date: _________________________________",
        ]
        for line in bcc_lines:
            c0.add_paragraph(line)
        for line in client_lines:
            c1.add_paragraph(line)
        break

    # Tenth pass: collapse excessive consecutive empty paragraphs (max 2 in a row)
    _collapse_empty_paragraphs(doc, max_consecutive=2)

    set_all_text_black(doc)

    # Enable automatic field update on open (TOC page numbers stay in sync)
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        settings = doc.settings.element
        # Remove existing updateFields then add fresh one
        for el in settings.findall(qn("w:updateFields")):
            settings.remove(el)
        uf = OxmlElement("w:updateFields")
        uf.set(qn("w:val"), "true")
        settings.append(uf)
    except Exception as e:
        print(f"[warn] 无法设置 updateFields: {e}")

    doc.save(str(out_path))
    return True


def set_all_text_black(doc) -> None:
    """将文档中所有 run 的字体设为黑色，消除红色占位符。"""
    try:
        from docx.shared import RGBColor
    except ImportError:
        return
    black = RGBColor(0, 0, 0)
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.color.rgb = black
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = black
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            for p in header_footer.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = black


def save_proposal_md(project: dict, price_per_visit: int, est_visits: int) -> Path:
    """
    Generate a structured human-readable .md proposal draft for Telegram review.
    Saved to Pending_Approval/Outbound/Proposal_Draft_<ProjectName>.md.
    """
    outbound = BASE_DIR / "Pending_Approval" / "Outbound"
    outbound.mkdir(parents=True, exist_ok=True)

    proj_name = project.get("name", "Project")
    client = project.get("client", "Client")
    safe = sanitize_dirname(proj_name).replace(" ", "_")
    md_path = outbound / f"Proposal_Draft_{safe}.md"

    total = price_per_visit * est_visits
    desc = project.get("description") or (
        f"Tenant renovation / construction at {project.get('address', 'the project location')}."
    )
    today = datetime.now().strftime("%Y-%m-%d %H:%M")

    disciplines = "\n".join(
        f"- **{d}**: Applicable per permit set"
        for d in ["Building", "Mechanical", "Electrical", "Plumbing", "Fire Protection"]
    )
    extra_visits = max(0, est_visits - 4)

    content = f"""# Proposal Draft — {proj_name}

**Status**: PENDING REVIEW
**Generated**: {today}
**Source**: Automated pipeline

---

## Project Info

| Field | Value |
|-------|-------|
| **Project** | {proj_name} |
| **Address** | {project.get('address', 'N/A')} |
| **Client (GC)** | {client} |
| **Contact** | {project.get('attention', project.get('contact_name', 'N/A'))} |
| **Email** | {project.get('contact_email', 'N/A')} |

## Project Description

{desc}

## Proposed Scope of Work (BCC)

Third-Party Code Inspection Services for {proj_name}.

### Applicable Disciplines
{disciplines}

### Estimated Inspections

| Inspection Type | Visits |
|-----------------|--------|
| Rough-in (MEP, framing) | 1–2 |
| Close-in | 1 |
| Final inspection | 1 |
| Additional (if required) | {extra_visits} |
| **Total Estimated Visits** | **{est_visits}** |

### Fee Estimate

| Item | Value |
|------|-------|
| Price per visit | ${price_per_visit} |
| Estimated visits | {est_visits} |
| **Estimated Total** | **${total}** |

---

## Next Steps

1. **Review on Telegram**: Reply `OK` to approve, or send changes
2. **Generate**: Word doc + PDF generated after approval
3. **Email**: Outreach email prepared for {project.get('attention', project.get('contact_name', 'contact'))}

---

*Reply `OK` to approve | `price 350` to change fee | `visits 6` to change count*
"""
    md_path.write_text(content, encoding="utf-8")
    return md_path


def create_email_draft(client_name: str, project_name: str, proposal_path: Path, to_email: str = "") -> Path:
    """在 Pending_Approval/Outbound/ 下创建发送提案的邮件草稿。"""
    outbound = BASE_DIR / "Pending_Approval" / "Outbound"
    outbound.mkdir(parents=True, exist_ok=True)
    safe = sanitize_dirname(project_name).replace(" ", "_")
    draft_path = outbound / f"BC_Proposal_{safe}_Draft.md"
    body = f"""**收件人**：{to_email or '(请填写 BC 项目联系人邮箱)'}
**邮箱**：（请填写）
**Subject:** Third-Party Code Compliance Inspection Proposal – {project_name}

---

Building Code Consulting LLC 已根据项目信息准备好 DC 第三方检测服务提案，见附件。

Please find attached our proposal for Third-Party Code Compliance Inspection services for {project_name}. We are a DC Third-Party agency; Yue Cao (PE, MCP) will oversee code compliance and inspection coordination.

Best regards,
Yue Cao, PE, MCP
Building Code Consulting
"""
    draft_path.write_text(
        f"# 邮件草稿：BuildingConnected 提案 – {project_name}\n\n"
        f"**附件**：{proposal_path.name}\n\n"
        + body,
        encoding="utf-8",
    )
    return draft_path


def docx_to_pdf(docx_path: Path) -> Path | None:
    """
    将生成的 .docx 转为同目录下的 .pdf（Windows 下使用 docx2pdf/Word COM）。
    返回 PDF 路径，失败返回 None。
    """
    if not docx_path or not Path(docx_path).exists():
        return None
    docx_path = Path(docx_path)
    pdf_path = docx_path.with_suffix(".pdf")
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        return pdf_path if pdf_path.exists() else None
    except Exception as e:
        print(f"PDF 转换失败 {docx_path}: {e}", file=sys.stderr)
        return None


def run_single_project(
    project: dict,
    price_per_visit: int | None = None,
    est_visits: int = 12,
    skip_confirm: bool = False,
    template_path: Path | None = None,
    template_type: str = "code_compliance",
    telegram_review: bool = False,
) -> dict:
    """
    对单个项目：建议定价 → 可选 Telegram 审核 → 填模板 → 输出 Word + PDF → 写邮件草稿。
    template_type: "code_compliance" | "plan_review"；未传 template_path 时按此选模板。
    telegram_review=True: 先生成 .md 草稿，发送到 Telegram 等待批复，再生成 Word 文件。
    返回 { "success", "output_docx", "pdf_path", "md_draft_path", "draft_path", "error" }。
    """
    template_path = template_path or get_template_path(template_type)
    tier, suggested_price, reason = suggest_tier(project)
    price = price_per_visit if price_per_visit is not None else suggested_price

    # --- Step 1: Save structured .md draft ---
    md_draft_path = save_proposal_md(project, price, est_visits)
    print(f"[Proposal] .md draft saved: {md_draft_path}")

    # --- Step 2: Telegram review loop (if requested) ---
    if telegram_review:
        try:
            from core_tools.telegram_approval import run_telegram_approval_loop
            approved, project, price, est_visits = run_telegram_approval_loop(
                md_draft_path, project, price, est_visits,
                max_rounds=3, timeout_per_round=120,
            )
            if not approved:
                return {
                    "success": False,
                    "output_docx": "",
                    "pdf_path": "",
                    "md_draft_path": str(md_draft_path),
                    "draft_path": "",
                    "error": f"Telegram approval not received — proposal saved to {md_draft_path}",
                }
        except Exception as e:
            print(f"[Warning] Telegram review failed ({e}). Falling back to terminal confirm.", file=sys.stderr)
            if not skip_confirm:
                print_pricing_summary(project, tier, price, reason, est_visits)
    elif not skip_confirm:
        print_pricing_summary(project, tier, price, reason, est_visits)

    # --- Step 3: Generate Word doc ---
    out_dir = get_proposal_output_dir(project.get("client", "Unknown"), project.get("name", "Project"))
    docx_name = f"{sanitize_dirname(project.get('name', 'Proposal'))} - Third Party Code Inspection Proposal from BCC.docx"
    out_docx = out_dir / docx_name
    scope = (project.get("description") or "")[:3000]
    ok = fill_template(template_path, out_docx, project, price, est_visits, scope)
    if not ok:
        return {
            "success": False, "output_docx": "", "pdf_path": "",
            "md_draft_path": str(md_draft_path), "draft_path": "",
            "error": "模板填充失败",
        }

    # --- Step 4: Generate PDF ---
    pdf_path = docx_to_pdf(out_docx)
    if pdf_path:
        print(f"[Proposal] PDF generated: {pdf_path}")

    # --- Step 5: Create email draft ---
    draft_path = create_email_draft(
        project.get("client", ""),
        project.get("name", ""),
        out_docx,
        project.get("contact_email", ""),
    )

    # --- Step 6: Notify via Telegram (if review was done there) ---
    if telegram_review:
        try:
            from core_tools.telegram_approval import send_message
            notify = f"📄 Word doc ready:\n`{out_docx.name}`"
            if pdf_path:
                notify += f"\n📎 PDF also generated: `{Path(pdf_path).name}`"
            notify += f"\n\n📁 Saved to:\n`{out_dir}`"
            send_message(notify)
        except Exception:
            pass

    # --- Step 7: Log to work_log ---
    try:
        from core_tools.work_log import mark_proposal_done, mark_email_drafted
        mark_proposal_done(project.get("client", ""), project.get("name", ""), str(out_docx))
        mark_email_drafted(project.get("client", ""), project.get("name", ""), str(draft_path))
    except Exception:
        pass

    return {
        "success": True,
        "output_docx": str(out_docx),
        "pdf_path": str(pdf_path) if pdf_path else "",
        "md_draft_path": str(md_draft_path),
        "draft_path": str(draft_path),
        "error": "",
    }


def main():
    import argparse
    ap = argparse.ArgumentParser(description="DC 第三方检测/Plan Review 提案生成（Word 模板 + 定价逻辑）")
    ap.add_argument("--client", default="", help="客户名称")
    ap.add_argument("--project", default="", help="项目名称")
    ap.add_argument("--price", type=int, default=None, help="指定单价 $/visit，不填则自动建议")
    ap.add_argument("--visits", type=int, default=12, help="预估检测次数")
    ap.add_argument("--skip-confirm", action="store_true", help="不打印定价表，直接生成")
    ap.add_argument("--template", type=str, default="", help="Word 模板路径（不填则用 --type 选择）")
    ap.add_argument("--type", dest="template_type", choices=("code_compliance", "plan_review"), default="code_compliance", help="使用 Code Compliance 或 Plan Review 模板")
    ap.add_argument("--contact", default="", help="Contact/attention person name")
    ap.add_argument("--address", default="", help="Project address (e.g. '20 F Street NW, Suite 550, Washington, DC 20001')")
    ap.add_argument("--email", default="", help="Contact email address")
    ap.add_argument("--description", default="", help="Project description / scope notes")
    ap.add_argument("--disciplines", default="building,mechanical,electrical,plumbing,fire_protection",
                    help="Comma-separated applicable disciplines. Default: all 5. "
                         "Valid values: building, mechanical, electrical, plumbing, fire_protection")
    ap.add_argument("--telegram-review", action="store_true", help="Send .md draft to Telegram, wait for approval before generating Word doc")
    args = ap.parse_args()

    # Parse disciplines into a normalized set
    disciplines_set = {
        d.strip().lower().replace(" ", "_").replace("-", "_")
        for d in (args.disciplines or "").split(",")
        if d.strip()
    }
    # Map common aliases
    _DISC_ALIASES = {"fp": "fire_protection", "fire": "fire_protection",
                     "mech": "mechanical", "elec": "electrical", "plb": "plumbing",
                     "bldg": "building", "bl": "building"}
    disciplines_set = {_DISC_ALIASES.get(d, d) for d in disciplines_set}

    project = {
        "name": args.project or "St. Joseph's on Capitol Hill – Phase I",
        "client": args.client or "Sample Client",
        "attention": args.contact,
        "address": args.address,
        "contact_email": args.email,
        "description": args.description,
        "disciplines": disciplines_set,
    }
    template_path = Path(args.template) if args.template else None
    result = run_single_project(
        project,
        price_per_visit=args.price,
        est_visits=args.visits,
        skip_confirm=args.skip_confirm,
        template_path=template_path,
        template_type=args.template_type,
        telegram_review=args.telegram_review,
    )
    if result["success"]:
        print("提案已生成:", result["output_docx"])
        if result.get("pdf_path"):
            print("PDF:", result["pdf_path"])
        print(".md 草稿:", result.get("md_draft_path", ""))
        print("邮件草稿:", result["draft_path"])
    else:
        print("失败:", result.get("error"), file=sys.stderr)
    sys.exit(0 if result["success"] else 1)


if __name__ == "__main__":
    main()
