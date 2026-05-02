"""Generate a single-page residential SFR pricing sheet for BCC partners/clients.

Produces BCC_Residential_SFR_Pricing.docx in this folder. Two-option layout:
  Option A — Flat Rate per Visit ($350)
  Option B — Per-Inspection Schedule (single-trade $290, combo $550 / $800)
"""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.parser import parse_xml
from docx.shared import Pt, Inches, RGBColor


OUT = Path(__file__).parent / "BCC_Residential_SFR_Pricing.docx"

OFFICE_PHONE = "571-365-6937 / 806-787-5806"
OFFICE_ADDR = "3914 Tallow Tree Ct, Fairfax, VA 22033"

SINGLE_TRADE = [
    ("Footing", "$290"),
    ("Foundation", "$290"),
    ("Slab", "$290"),
    ("Groundworks", "$290"),
    ("Framing / Sheathing / Bracing", "$290"),
    ("TPF", "$290"),
    ("Insulation (per permit)", "$290"),
    ("Blower Door Test (per test)", "$290"),
    ("Duct Pressure Test (per test)", "$290"),
]

COMBO = [
    ("MEP Roughs", "$550 (≤4 trades)  /  $800 (5+ trades)"),
    ("Finals", "$550 (≤4 trades)  /  $800 (5+ trades)"),
]

ADD_ON = [
    ("Building Envelope (when adjacent to another visit)", "$50"),
    ("Close-In, drywall hung no mud/tape (when adjacent)", "$100"),
]

SPECIALTY = [
    ("DOB Abatement Inspection", "$350"),
    ("Hourly / Re-inspection Rate", "$290 / hr"),
]

LOGO_PATH = (
    Path("C:/Users/Kyle Cao/DC Business/Building Code Consulting/")
    / "Logo E-Sig Stamp" / "BCC Logo New.png"
)

EMU_PER_INCH = 914400


def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)


def add_para(doc, text, size=10, bold=False, italic=False, align=None, space_after=2):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_section_header(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_price_table(doc, rows, col_widths=(Inches(5.0), Inches(1.2))):
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    for row_idx, (label, price) in enumerate(rows):
        cells = table.rows[row_idx].cells
        cells[0].width = col_widths[0]
        cells[1].width = col_widths[1]
        for cell, txt, bold in ((cells[0], label, False), (cells[1], price, True)):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(txt)
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            run.bold = bold
            run.font.color.rgb = RGBColor(0, 0, 0)
            if txt == price:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return table


def floatify_picture(run, pos_h_in, pos_v_in, h_rel="page", v_rel="page",
                     z_index=251658240):
    """Convert the inline picture in `run` to a floating anchored picture
    positioned absolutely on the page. wrapNone — text flows independently."""
    drawing = run._element.find(qn("w:drawing"))
    inline = drawing.find(qn("wp:inline"))

    extent = inline.find(qn("wp:extent"))
    docPr = inline.find(qn("wp:docPr"))
    cNv = inline.find(qn("wp:cNvGraphicFramePr"))
    graphic = inline.find(qn("a:graphic"))

    cx = extent.get("cx")
    cy = extent.get("cy")
    pos_h_emu = int(pos_h_in * EMU_PER_INCH)
    pos_v_emu = int(pos_v_in * EMU_PER_INCH)

    nsdecl = (
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
    )
    anchor_xml = (
        f'<wp:anchor {nsdecl} '
        f'distT="0" distB="0" distL="114300" distR="114300" simplePos="0" '
        f'relativeHeight="{z_index}" behindDoc="0" locked="0" '
        f'layoutInCell="1" allowOverlap="1">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="{h_rel}"><wp:posOffset>{pos_h_emu}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="{v_rel}"><wp:posOffset>{pos_v_emu}</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:wrapNone/>'
        f'</wp:anchor>'
    )
    anchor = parse_xml(anchor_xml)
    anchor.append(docPr)
    anchor.append(cNv)
    anchor.append(graphic)

    drawing.remove(inline)
    drawing.append(anchor)


def add_paragraph_bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def build():
    doc = Document()
    set_margins(doc)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # === LETTERHEAD ===
    # Centered company name across full page width
    title_p = add_para(
        doc, "BUILDING CODE CONSULTING LLC",
        size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0,
    )

    # Float the logo at top-left of page, anchored to the title paragraph.
    # The logo sits in the top-left corner; the centered title text passes
    # well to the right of it without colliding.
    if LOGO_PATH.exists():
        logo_run = title_p.add_run()
        logo_run.add_picture(str(LOGO_PATH), width=Inches(1.0))
        floatify_picture(
            logo_run,
            pos_h_in=0.45, pos_v_in=0.35,
            h_rel="page", v_rel="page",
        )

    add_para(
        doc,
        "Third-Party Inspection Agency  |  Plan Review  |  Code Consulting",
        size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0,
    )
    add_para(
        doc,
        "admin@buildingcodeconsulting.com  |  ycao@buildingcodeconsulting.com",
        size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0,
    )
    rule_p = add_para(
        doc,
        f"{OFFICE_PHONE}   |   {OFFICE_ADDR}",
        size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4,
    )
    add_paragraph_bottom_border(rule_p)

    # === TITLE ===
    add_para(
        doc, "Residential — Third Party Inspection Pricing",
        size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0,
    )
    add_para(
        doc, f"Updated {date.today().strftime('%B %Y')}",
        size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4,
    )

    # === INTRO ===
    add_para(
        doc,
        "Thank you for the opportunity to provide pricing for Third Party Building "
        "Inspection services on Single Family Renovation and New Construction projects. "
        "BCC offers two pricing options — choose the one that best fits your project.",
        size=10, space_after=4,
    )

    # === OPTION A ===
    add_section_header(doc, "OPTION A — FLAT RATE PER VISIT:  $350 / visit")
    add_para(
        doc,
        "One simple price for any required inspection type. Billed per actual visit "
        "completed — no upfront estimation required. Best for small renovations and "
        "projects without a finalized drawing set.",
        size=10, space_after=4,
    )

    # === OPTION B ===
    add_section_header(doc, "OPTION B — PER-INSPECTION SCHEDULE")

    add_para(doc, "Single-Trade / Sequential Inspections", size=10, bold=True, space_after=1)
    add_price_table(doc, SINGLE_TRADE)

    add_para(doc, "Multi-Trade Combo Inspections", size=10, bold=True, space_after=1)
    add_price_table(doc, COMBO, col_widths=(Inches(3.4), Inches(2.8)))

    add_para(doc, "Add-On / Quick Verifications", size=10, bold=True, space_after=1)
    add_price_table(doc, ADD_ON)

    add_para(doc, "Specialty", size=10, bold=True, space_after=1)
    add_price_table(doc, SPECIALTY)

    # === TERMS ===
    add_section_header(doc, "SCHEDULING & TERMS")
    terms = [
        "Same-day or next-business-day inspection scheduling.",
        "No after-hours, weekend, or fuel surcharge — same price any time, any day.",
        "Permits and paperwork required 2 business days before the first inspection "
        "so we can secure third-party approval through the local AHJ.",
        "Subsequent inspections scheduled the next business day if requested by 1pm the prior business day.",
        "Reports filed with the city the next business day after inspection (PE QA review included).",
        "Cancellations after 5pm the business day before scheduled inspection charged full price.",
        "Invoices submitted weekly; payment due within 15 days. "
        "Accounts 45+ days past due placed on credit hold.",
    ]
    for t in terms:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(t)
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    # === CONTACT ===
    add_section_header(doc, "CONTACT")
    add_para(doc, "Yue Cao, PE, MCP — President", size=10, bold=True, space_after=0)
    add_para(doc, f"{OFFICE_PHONE}   |   {OFFICE_ADDR}", size=10, space_after=0)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
