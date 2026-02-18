"""Inspect .docx for all text (paragraphs, tables, headers, footers) to find where wrong content lives."""
from pathlib import Path
from docx import Document

path = Path(r"c:\Users\Kyle Cao\DC Business\Building Code Consulting\Projects\Keller Brothers\St. Joseph's on Capitol Hill – Phase I\St. Joseph's on Capitol Hill – Phase I - Third Party Code Inspection Proposal from BCC.docx")
if not path.exists():
    path = Path(r"c:\Users\Kyle Cao\DC Business\Building Code Consulting\Business Automation\DC Code Compliance Proposal Template.docx")
if not path.exists():
    path = Path(r"c:\Users\Kyle Cao\DC Business\Building Code Consulting\DC Code Compliance Proposal Template.docx")

doc = Document(str(path))
print("=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        print(f"[{i}] {t[:120]}{'...' if len(t) > 120 else ''}")
print("\n=== TABLE CELLS ===")
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            t = cell.text.strip()
            if t:
                print(f"T{ti}R{ri}C{ci}: {t[:100]}{'...' if len(t) > 100 else ''}")
print("\n=== HEADERS / FOOTERS ===")
for sec in doc.sections:
    for part_name, part in [("header", sec.header), ("footer", sec.footer)]:
        for p in part.paragraphs:
            t = p.text.strip()
            if t:
                print(f"{part_name}: {t[:100]}")
print("\n=== SEARCH: Insomnia, Monroe, 01-12, Cox, Bryan, United States ===")
def search_in_para(p, prefix=""):
    for run in p.runs:
        t = run.text
        for kw in ["Insomnia", "Monroe", "01-12", "Cox", "Bryan", "United States", "tenant fit out", "AHUs"]:
            if kw in t:
                print(f"FOUND '{kw}' in {prefix}: {repr(t[:80])}")
for i, p in enumerate(doc.paragraphs):
    search_in_para(p, f"para[{i}]")
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            for pi, p in enumerate(cell.paragraphs):
                search_in_para(p, f"T{ti}R{ri}C{ci}p{pi}")
for sec in doc.sections:
    for p in sec.header.paragraphs:
        search_in_para(p, "header")
    for p in sec.footer.paragraphs:
        search_in_para(p, "footer")
