"""
Generate a BCC Master Inspection Contract (MIC) from the canonical template.

Template:  Master Contract And ATP/BCC Third Party Code Compliance Inspection Service Master Proposal Template.docx
Output:    Master Contract And ATP/<Client>/BCC Third Party Code Compliance Inspection Service Master Proposal for <Client>.docx

When called for a master proposal / master contract / MIC, USE THIS — not the
project-style proposal template. The MIC is a real one-year auto-renewing
contract with the ATP per-project mechanism baked in (Article 1.1).
"""
import argparse
import re
import sys
import copy
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("pip install python-docx")
    sys.exit(1)

ROOT = Path(r"C:\Users\Kyle Cao\DC Business\Building Code Consulting\Master Contract And ATP")
TEMPLATE = ROOT / "BCC Third Party Code Compliance Inspection Service Master Proposal Template.docx"


def next_mic_number() -> str:
    """Scan existing MICs in Master Contract And ATP/ and return next sequential number."""
    used = set()
    for f in ROOT.rglob("*.docx"):
        try:
            d = Document(str(f))
            for p in d.paragraphs[:15]:
                m = re.search(r"MIC-(\d{4})-(\d{3})", p.text)
                if m and f.name != TEMPLATE.name:
                    used.add(int(m.group(2)))
                    break
        except Exception:
            pass
    n = (max(used) + 1) if used else 1
    year = datetime.today().year
    return f"MIC-{year}-{n:03d}"


def _set_black(run):
    run.font.color.rgb = RGBColor(0, 0, 0)


def _replace_in_runs(p, old: str, new: str) -> bool:
    """Find old in concatenated run text and replace, preserving first-run format. Returns True if replaced."""
    full = "".join(r.text for r in p.runs)
    if old not in full:
        return False
    # Capture first run formatting
    if p.runs:
        r0 = p.runs[0]
        fmt = (r0.bold, r0.italic, r0.font.size, r0.font.name, r0.underline)
    else:
        fmt = (None, None, None, None, None)
    new_full = full.replace(old, new)
    p.clear()
    run = p.add_run(new_full)
    run.bold, run.italic = fmt[0], fmt[1]
    if fmt[2] is not None: run.font.size = fmt[2]
    if fmt[3] is not None: run.font.name = fmt[3]
    if fmt[4] is not None: run.underline = fmt[4]
    _set_black(run)
    return True


def _insert_para_after(ref_p, text: str):
    """Insert a new paragraph immediately after ref_p, copying its style."""
    new_p_xml = OxmlElement("w:p")
    # Copy paragraph properties (style/indent) if present
    pPr = ref_p._p.find(qn("w:pPr"))
    if pPr is not None:
        new_p_xml.append(copy.deepcopy(pPr))
    ref_p._p.addnext(new_p_xml)
    # Wrap as Paragraph
    from docx.text.paragraph import Paragraph
    new_p = Paragraph(new_p_xml, ref_p._parent)
    run = new_p.add_run(text)
    if ref_p.runs:
        r0 = ref_p.runs[0]
        if r0.font.size is not None: run.font.size = r0.font.size
        if r0.font.name is not None: run.font.name = r0.font.name
    _set_black(run)
    return new_p


def set_all_text_black(doc):
    black = RGBColor(0, 0, 0)
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.color.rgb = black
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = black


def generate(client_name: str, single_rate: int = 325, combo_rate: int | None = None,
             hourly_rate: int = 150, mic_no: str | None = None, out_filename: str | None = None,
             effective_date: str | None = None) -> Path:
    """
    client_name      — e.g. "Team VP Construction"
    single_rate      — $/single-trade visit (default 325)
    combo_rate       — $/combination visit. If None, omit second tier (single rate is the only rate).
    hourly_rate      — additional consultation $/hour (default 150)
    mic_no           — override MIC reference; default = next sequential
    out_filename     — override output filename; default = "BCC Third Party Code Compliance Inspection Service Master Proposal for <Client>.docx"
    effective_date   — string (e.g. "04-30-2026"); default = today MM-DD-YYYY
    """
    if not TEMPLATE.exists():
        print(f"[ERR] Template not found: {TEMPLATE}")
        sys.exit(1)

    mic_no = mic_no or next_mic_number()
    effective_date = effective_date or datetime.today().strftime("%m-%d-%Y")
    today_sig = datetime.today().strftime("%-m/%-d/%Y") if sys.platform != "win32" else datetime.today().strftime("%#m/%#d/%Y")

    out_dir = ROOT / client_name
    out_dir.mkdir(parents=True, exist_ok=True)
    out_file = out_dir / (out_filename or f"BCC Third Party Code Compliance Inspection Service Master Proposal for {client_name}.docx")

    d = Document(str(TEMPLATE))

    # 1. Replace MIC reference number, client name placeholder, effective date, signature date
    # 2. Find Article 4 "Standard Visit Rate:" paragraph and replace with two-tier if combo_rate provided
    art4_p = None
    sig_date_p = None
    for p in d.paragraphs:
        _replace_in_runs(p, "MIC-2026-005", mic_no)
        _replace_in_runs(p, "[Client Company Name]", client_name)
        _replace_in_runs(p, "[填入签署日期]", effective_date)
        # Signature date for BCC side ("Date:       3/24/2026")
        if p.text.strip().startswith("Date:") and "3/24/2026" in p.text:
            _replace_in_runs(p, "3/24/2026", today_sig)

        if p.text.strip().startswith("Standard Visit Rate:"):
            art4_p = p

    # Two-tier rate substitution in Article 4
    if art4_p is not None:
        if combo_rate is None:
            # Single-rate — just substitute the dollar amount
            _replace_in_runs(art4_p, "$375.00", f"${single_rate}.00")
        else:
            # Two-tier — rewrite this paragraph and insert a second one after it
            _replace_in_runs(art4_p,
                "Standard Visit Rate: $375.00 per inspection visit (Includes up to 2 hours on-site and 1 hour travel time).",
                f"Single-Trade Visit Rate: ${single_rate}.00 per inspection visit — one discipline per visit "
                f"(Building, Mechanical, Electrical, Plumbing, or Fire Protection). Includes up to 2 hours on-site and 1 hour travel time."
            )
            _insert_para_after(art4_p,
                f"Combination Visit Rate: ${combo_rate}.00 per inspection visit — two or more disciplines covered "
                f"in a single site visit. Includes up to 2 hours on-site and 1 hour travel time."
            )

    # Update hourly rate if non-default
    if hourly_rate != 150:
        for p in d.paragraphs:
            if p.text.strip().startswith("Additional Hourly Rate:"):
                _replace_in_runs(p, "$150.00", f"${hourly_rate}.00")

    set_all_text_black(d)
    d.save(str(out_file))
    print(f"[OK] {mic_no}  →  {out_file}")
    return out_file


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--client", required=True, help='e.g. "Team VP Construction"')
    ap.add_argument("--single-rate", type=int, default=325)
    ap.add_argument("--combo-rate", type=int, default=None,
                    help="If set, adds a second-tier combination visit rate. Omit for single-rate MIC.")
    ap.add_argument("--hourly-rate", type=int, default=150)
    ap.add_argument("--mic-no", default=None, help="Override; default = next sequential")
    ap.add_argument("--effective-date", default=None, help="MM-DD-YYYY; default = today")
    args = ap.parse_args()
    generate(args.client, args.single_rate, args.combo_rate,
             args.hourly_rate, args.mic_no, effective_date=args.effective_date)
