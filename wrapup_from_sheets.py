"""One-shot Plan Review Wrap-Up bootstrap.

Usage:
    python wrapup_from_sheets.py <google_sheet_url> [<google_sheet_url> ...]
    python wrapup_from_sheets.py <url> --send         # auto-send after self-review (still requires SMTP success)

Trigger pattern: Kyle pastes 1+ Google Spreadsheet URLs (one per discipline)
and says "wrap up". This script does the full chain:

  URLs → file IDs → download sheets → read project address →
  locate project folder → locate approved NOI → extract TPR # + dates →
  generate Docs #1-#5 → flatten ALL → self-review → save draft email →
  WAIT for Kyle's "Y" (script prints preview and exits unless --send is on AND
  preflight passes) → SMTP send + rename draft to -SENT.md.

Per BCC critical rule: --send only sends if every preflight check passes
(every PDF exists, has real content, has 0 widgets, recipient resolved).
Without --send the script always halts at the preview stage.

This file is the single entry point — keep its surface stable. Helper modules
(generate_*, flatten_*, send_*) it imports remain independently re-runnable.
"""
from __future__ import annotations
import argparse
import re
import shutil
import sys
from copy import deepcopy
from pathlib import Path

import fitz

REPO = Path(r"C:\Users\Kyle Cao\DC Business\Building Code Consulting\Business Automation")
PROJECTS_ROOT = Path(r"C:\Users\Kyle Cao\DC Business\Building Code Consulting\Projects")
TPL_DIR = Path(r"C:\Users\Kyle Cao\DC Business\Building Code Consulting\PLAN REVIEW PROJECT DOC TEMPLATE")
WRAP_TPL_DIR = REPO / "Wrap up"
GUIDE_MANUAL = TPL_DIR / "Third-Party_Program_Procedure_Manual 5.15.2023 seal.pdf"
APPROVAL_CERT_PAGE = 60  # 0-indexed page 61 of guide manual
TRACKER = REPO / "BCC_Outreach_Tracker.md"
PENDING = REPO / "Pending_Approval" / "Outbound"

DISCIPLINE_BLANK_TABLE = {
    frozenset({"FA", "SPK"}): "Fire(electrical) & Sprinkler",
    frozenset({"FA"}): "Fire Alarm",
    frozenset({"SPK"}): "Sprinkler",
}
LETTER_DISCIPLINES_TABLE = {
    frozenset({"FA", "SPK"}): ["Fire Protection System", "Sprinkler System", "Fire Alarm System"],
    frozenset({"FA"}): ["Fire Protection System", "Fire Alarm System"],
    frozenset({"SPK"}): ["Fire Protection System", "Sprinkler System"],
}
LARGESTAMP_ROWS_TABLE = {
    frozenset({"FA", "SPK"}): ["Fire Protection", "Sprinkler System"],
    frozenset({"FA"}): ["Fire Protection"],
    frozenset({"SPK"}): ["Sprinkler System"],
}
DOC3_DISC_ROWS_TABLE = {
    frozenset({"FA", "SPK"}): ["Fire Protection", "Sprinkler System"],
    frozenset({"FA"}): ["Fire Protection"],
    frozenset({"SPK"}): ["Sprinkler System"],
}


def fail(msg: str, *, code: int = 2) -> None:
    print(f"\n[HALT] {msg}", file=sys.stderr)
    sys.exit(code)


# ===== Step 1: parse URLs → file IDs =====
def extract_file_id(url: str) -> str:
    m = re.search(r"/d/([a-zA-Z0-9_-]{20,})", url)
    if not m:
        fail(f"Could not extract file ID from URL: {url}")
    return m.group(1)


# ===== Step 2: download via Drive MCP wrapper =====
def download_sheet_via_mcp(file_id: str, dst_xlsx: Path) -> Path:
    """The Drive MCP returns a JSON wrapper {content:[{embeddedResource:{contents:{blob:<base64>}}}]}.
    This script can't itself call MCP from inside a normal python invocation,
    so it expects an already-saved JSON tool-result file at:
        <REPO>/_drive_cache/<file_id>.json
    Claude (the agent) is responsible for invoking the Drive MCP and dropping
    the tool-result file into that path before running this step.
    """
    cache = REPO / "_drive_cache"
    cache.mkdir(exist_ok=True)
    json_in = cache / f"{file_id}.json"
    if not json_in.exists():
        fail(f"Drive MCP cache miss for {file_id}.\n"
             f"Claude should call mcp__claude_ai_Google_Drive__download_file_content\n"
             f"  fileId={file_id}\n"
             f"  exportMimeType=application/vnd.openxmlformats-officedocument.spreadsheetml.sheet\n"
             f"and write the JSON tool-result to {json_in}")
    import json, base64
    payload = json.loads(json_in.read_text(encoding="utf-8"))
    blob = payload["content"][0]["embeddedResource"]["contents"]["blob"]
    dst_xlsx.write_bytes(base64.b64decode(blob))
    return dst_xlsx


# ===== Step 3: read sheet metadata =====
def read_sheet_metadata(xlsx_path: Path) -> dict:
    """Use openpyxl (read-only, no Excel lock). Returns:
        {project_address, discipline_tag (FA|SPK|UNKNOWN), filled_last_row, filled_last_col}
    Project address is the BCC deficiency template's row 3 col D, which holds
    e.g. "1522 Rhode Island Ave NE". Discipline is inferred from the sheet's
    column headers / file name."""
    from openpyxl import load_workbook
    wb = load_workbook(xlsx_path, read_only=True, data_only=True)
    ws = wb.active
    # Address typically lives at row 3, col D in the BCC template
    addr = ws.cell(row=3, column=4).value
    # Discipline detection: look at sheet title + nearby cells for SPK/FA tokens
    haystack = " ".join([ws.title or ""] + [str(ws.cell(row=r, column=c).value or "")
                                            for r in range(1, 13) for c in range(2, 8)])
    tag = "UNKNOWN"
    h_up = haystack.upper()
    if "FIRE ALARM" in h_up or " FA " in f" {h_up} " or h_up.startswith("FA "):
        tag = "FA"
    if "SPRINKLER" in h_up or "SPK" in h_up:
        # Either SPK or both. Prefer FA if it clearly says fire alarm too.
        if tag != "FA":
            tag = "SPK"
    # Filled-row detection (col B = #, col A empty padding; real content cols 3-7)
    filled_first_row = 1
    filled_last_row = filled_first_row + 3  # at minimum keep the header band
    last_row = ws.max_row or 1
    last_col = ws.max_column or 7
    for r in range(last_row, 4, -1):
        if any(_nonempty(ws.cell(row=r, column=c).value) for c in range(3, last_col + 1)):
            filled_last_row = r
            break
    wb.close()
    if not addr or not str(addr).strip():
        fail(f"Project address not found at expected cell D3 in {xlsx_path.name}.")
    return {
        "project_address": str(addr).strip(),
        "discipline_tag": tag,
        "filled_first_row": filled_first_row,
        "filled_last_row": filled_last_row,
        "filled_last_col": last_col,
    }


def _nonempty(v) -> bool:
    if v is None: return False
    if isinstance(v, str) and not v.strip(): return False
    return True


# ===== Step 4: locate project folder =====
def locate_project_folder(addr: str) -> Path:
    """Project folders nest as Projects/<Client>/<Project ... addr ...>/.
    We search for a folder whose name contains the address minus filler tokens."""
    addr_norm = re.sub(r"[^A-Za-z0-9 ]+", " ", addr).strip()
    needles = [t for t in addr_norm.split() if len(t) >= 3 and not t.isdigit()] + \
              [t for t in addr_norm.split() if t.isdigit()]
    candidates = []
    for client_dir in PROJECTS_ROOT.iterdir():
        if not client_dir.is_dir(): continue
        for proj_dir in client_dir.iterdir():
            if not proj_dir.is_dir(): continue
            name = proj_dir.name.lower()
            score = sum(1 for n in needles if n.lower() in name)
            if score:
                candidates.append((score, proj_dir))
    if not candidates:
        fail(f"Could not locate a project folder under {PROJECTS_ROOT} for address: {addr}")
    candidates.sort(reverse=True, key=lambda t: (t[0], t[1].stat().st_mtime))
    return candidates[0][1]


# ===== Step 5: locate approved NOI =====
def locate_approved_noi(project_folder: Path) -> Path:
    matches = list(project_folder.glob("Approved_*NOI*signed*.pdf")) + \
              list(project_folder.glob("Approved_*NOI*.pdf"))
    matches = [m for m in matches if m.is_file()]
    if not matches:
        fail(f"No 'Approved_*NOI*signed*.pdf' in {project_folder}.\n"
             f"Either DOB hasn't stamped + returned the NOI yet, or it's saved under a different name.")
    matches.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    return matches[0]


# ===== Step 6: parse approved NOI =====
def parse_approved_noi(noi_pdf: Path) -> dict:
    """Page 1 Section A: project name + address. Page 2 FOR-OFFICIAL-USE-ONLY:
    DOB acceptance DATE and NOTICE OF APPROVAL CERTIFICATION NUMBER (TPR…)."""
    d = fitz.open(noi_pdf)
    p1 = d[0].get_text()
    p2 = d[1].get_text() if len(d) > 1 else ""
    d.close()

    # Page 1: project name (Section A.2) — the line after "PROJECT NAME"
    project_name = _extract_after(p1, r"PROJECT NAME[^\n]*", strip_chars=" _") or ""
    project_address = _extract_after(p1, r"PROJECT ADDRESS[^\n]*", strip_chars=" _") or ""

    # Page 2: TPR + DATE in the FOR-OFFICIAL-USE-ONLY box.
    # The form's filled values often appear on their own lines after the static labels.
    tpr_match = re.search(r"(TPR\s*\d{4,}[\-\d]*)", p2)
    if not tpr_match:
        fail(f"Could not find TPR approval number in {noi_pdf.name}. "
             f"Verify DOB has stamped the NOI (FOR-OFFICIAL-USE-ONLY box filled).")
    tpr = re.sub(r"\s+", "", tpr_match.group(1))

    date_match = re.search(r"(\d{1,2}/\d{1,2}/\d{2,4})", p2)
    if not date_match:
        fail(f"Could not find DOB acceptance date in {noi_pdf.name}.")
    date_raw = date_match.group(1)
    # Normalize to MM-DD-YYYY
    parts = re.split(r"[/-]", date_raw)
    if len(parts[2]) == 2:
        parts[2] = "20" + parts[2]
    deficiency_date = f"{int(parts[0]):02d}-{int(parts[1]):02d}-{parts[2]}"

    return {
        "project_name": project_name.strip() or project_address.strip(),
        "project_address_short": project_address.strip(),
        "tpr": tpr,
        "deficiency_date": deficiency_date,
    }


def _extract_after(text: str, label_pattern: str, strip_chars: str = " ") -> str | None:
    m = re.search(label_pattern, text)
    if not m: return None
    after = text[m.end():m.end() + 200]
    # Take up to the next newline run
    line = after.split("\n", 1)[0]
    return line.strip(strip_chars)


# ===== Step 9: recipient lookup =====
def lookup_recipient(project_address: str) -> tuple[str, str] | None:
    """Returns (name, email) by grepping BCC_Outreach_Tracker.md for the address.
    Tracker rows are pipe-delimited: | status | project name+addr | email | date | notes |"""
    if not TRACKER.exists():
        return None
    text = TRACKER.read_text(encoding="utf-8")
    # Each lead row contains the project address somewhere in cells 2 or 3.
    addr_low = project_address.lower()
    for line in text.splitlines():
        if not line.startswith("|"):
            continue
        if addr_low in line.lower():
            parts = [p.strip() for p in line.split("|")]
            email_cell = next((p for p in parts if "@" in p and "." in p), None)
            if email_cell:
                return ("(see tracker)", email_cell)
    return None


# ===== Step 10a: Doc #2 deficiency PDF print =====
def print_xlsx_to_pdf(xlsx_path: Path, pdf_out: Path, filled_last_row: int, filled_last_col: int) -> Path:
    import win32com.client as win32
    excel = win32.gencache.EnsureDispatch("Excel.Application")
    excel.Visible = False; excel.DisplayAlerts = False
    try:
        wb = excel.Workbooks.Open(str(xlsx_path), ReadOnly=True)
        try:
            ws = wb.Worksheets(1)
            rng = ws.Range(ws.Cells(1, 1), ws.Cells(filled_last_row, filled_last_col))
            ws.PageSetup.PrintArea = rng.Address
            ws.PageSetup.Orientation = 2
            ws.PageSetup.Zoom = False
            ws.PageSetup.FitToPagesWide = 1
            ws.PageSetup.FitToPagesTall = False
            for attr, val in [("LeftMargin", 0.25), ("RightMargin", 0.25),
                              ("TopMargin", 0.75), ("BottomMargin", 0.75),
                              ("HeaderMargin", 0.3), ("FooterMargin", 0.3)]:
                setattr(ws.PageSetup, attr, excel.InchesToPoints(val))
            wb.ExportAsFixedFormat(0, str(pdf_out))
        finally:
            wb.Close(SaveChanges=False)
    finally:
        excel.Quit()
    return pdf_out


# ===== Step 10b: Doc #3 AcroForm fill =====
def fill_doc3(out_path: Path, *, tpr: str, project_name: str, project_address: str,
              deficiency_date: str, today_date: str, day_month: str, year2: str,
              disciplines: list[str], disc_blank: str) -> Path:
    src = fitz.open(GUIDE_MANUAL)
    new = fitz.open()
    new.insert_pdf(src, from_page=APPROVAL_CERT_PAGE, to_page=APPROVAL_CERT_PAGE)
    src.close()
    page = new[0]
    fields = {
        "fill_1": tpr, "Date": today_date, "Permit Number": "",
        "Project Name_2": project_name, "Project Address_2": project_address,
        "day of": day_month, "20": year2,
        "Print Full Name and Title": "Yue Cao Professional in Charge",
        "ProfessionalinCharge of Third Party Plan Review Agency for": disc_blank,
        "Name of Agency": "Building Code Consulting",
        "Agency Approval ID Number": "TPR-05012025",
        "Professional EngineerArchitect or MCP Number": "PE920502",
    }
    # Discipline rows: row 1 → fill_15..17, row 2 → fill_19..21
    row_field_groups = [
        ("Plan Review DisciplineRow1", "fill_15", "fill_16", "fill_17"),
        ("Plan Review DisciplineRow2", "fill_19", "fill_20", "fill_21"),
    ]
    for i, disc in enumerate(disciplines[:2]):
        label_f, d1f, d2f, d3f = row_field_groups[i]
        fields[label_f] = disc
        fields[d1f] = deficiency_date
        fields[d2f] = today_date
        fields[d3f] = today_date
    for w in page.widgets() or []:
        if w.field_name in fields:
            w.field_value = fields[w.field_name]; w.update()
        elif w.field_name == "Fire_2":
            w.field_value = True; w.update()
    new.save(out_path, garbage=4, deflate=True)
    new.close()
    return out_path


# ===== Step 10c+d: Doc #4 + Doc #5 =====
def build_doc4(out_path: Path, project_address_full: str, today_date: str, disciplines: list[str]) -> Path:
    from docx import Document
    from docx.oxml.ns import qn
    src = WRAP_TPL_DIR / "1345 Connecticut Ave NW Fire Alarm Plan Review Certification Letter.docx"
    shutil.copy2(src, out_path)
    d = Document(out_path)
    mapping = {
        "12-09-2025": today_date,
        "1345 Connecticut Ave NW Washington DC": project_address_full,
        "Fire Protection System": "::DISC::",
    }
    _replace_text_runs(d, mapping)
    for p in list(d.paragraphs):
        if p.text.strip() == "::DISC::":
            p.runs[0].text = disciplines[0]
            for r in p.runs[1:]: r.text = ""
            xml_parent = p._p.getparent()
            idx = list(xml_parent).index(p._p)
            for j, disc in enumerate(disciplines[1:], start=1):
                new_p = deepcopy(p._p)
                xml_parent.insert(idx + j, new_p)
                ts = new_p.findall('.//' + qn('w:t'))
                if ts:
                    ts[0].text = disc
                    for t in ts[1:]: t.text = ""
            break
    d.save(out_path)
    return out_path


def build_doc5(out_path: Path, project_address_doc5: str, today_date_slash: str, disc_rows: list[str]) -> Path:
    from docx import Document
    from docx.oxml.ns import qn
    import re as _re
    src = WRAP_TPL_DIR / "Large Stamp Template.docx"
    shutil.copy2(src, out_path)
    d = Document(out_path)
    _replace_text_runs(d, {"1812 H PL SE, Washington D.C.": project_address_doc5})
    DRAWING = qn("w:drawing"); PICT = qn("w:pict")
    date_re = _re.compile(r"\d{1,2}[/-]\d{1,2}[/-]\d{2,4}")
    for tab in d.tables:
        for row in tab.rows:
            for cell in row.cells:
                for ntab in cell.tables:
                    for ri, nrow in enumerate(ntab.rows):
                        # Update label only if the template's existing row text doesn't match
                        if ri < len(disc_rows):
                            label_cell = nrow.cells[0]
                            cur = label_cell.text.strip()
                            if cur and cur != disc_rows[ri]:
                                # rewrite text-only runs (preserve any drawing runs)
                                for p in label_cell.paragraphs:
                                    full = "".join(r.text for r in p.runs)
                                    if cur in full:
                                        new_text = full.replace(cur, disc_rows[ri])
                                        placed = False
                                        for r in p.runs:
                                            if r._r.find(DRAWING) is not None or r._r.find(PICT) is not None:
                                                continue
                                            if not placed:
                                                r.text = new_text; placed = True
                                            else:
                                                r.text = ""
                                        break
                        date_cell = nrow.cells[2]
                        for p in date_cell.paragraphs:
                            full = "".join(r.text for r in p.runs)
                            if not date_re.search(full): continue
                            new_text = date_re.sub(today_date_slash, full, count=1)
                            placed = False
                            for r in p.runs:
                                if r._r.find(DRAWING) is not None or r._r.find(PICT) is not None:
                                    continue
                                if not placed:
                                    r.text = new_text; placed = True
                                else:
                                    r.text = ""
                            break
    d.save(out_path)
    return out_path


def _replace_text_runs(doc, mapping: dict) -> None:
    def visit(p):
        runs_text = "".join(r.text for r in p.runs)
        for old, new in mapping.items():
            if old in runs_text:
                runs_text = runs_text.replace(old, new)
                if p.runs:
                    p.runs[0].text = runs_text
                    for r in p.runs[1:]: r.text = ""
                return
    for p in doc.paragraphs:
        visit(p)
    for tab in doc.tables:
        for row in tab.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    visit(p)
                for ntab in cell.tables:
                    for nrow in ntab.rows:
                        for ncell in nrow.cells:
                            for np_ in ncell.paragraphs:
                                visit(np_)


# ===== Step 10e: Doc #4/#5 Word→PDF =====
def docx_to_pdf(docx: Path, pdf: Path) -> Path:
    import win32com.client as win32
    word = win32.gencache.EnsureDispatch("Word.Application")
    word.Visible = False; word.DisplayAlerts = 0
    try:
        doc = word.Documents.Open(str(docx), ReadOnly=True)
        try:
            doc.ExportAsFixedFormat(OutputFileName=str(pdf), ExportFormat=17,
                                    OpenAfterExport=False, OptimizeFor=0,
                                    BitmapMissingFonts=True, UseISO19005_1=False)
        finally:
            doc.Close(SaveChanges=False)
    finally:
        word.Quit()
    return pdf


# ===== Step 11: flatten =====
def flatten_pdf(src_path: Path, dpi: int = 250) -> None:
    src = fitz.open(src_path)
    out = fitz.open()
    for page in src:
        pix = page.get_pixmap(dpi=dpi, alpha=False)
        w_pt = pix.width * 72.0 / dpi; h_pt = pix.height * 72.0 / dpi
        new_page = out.new_page(width=w_pt, height=h_pt)
        new_page.insert_image(fitz.Rect(0, 0, w_pt, h_pt), pixmap=pix)
    src.close()
    tmp = src_path.with_suffix(".flat.pdf")
    out.save(tmp, garbage=4, deflate=True); out.close()
    src_path.unlink(); tmp.rename(src_path)


# ===== Step 12: self-review =====
def self_review_pdf(p: Path, *, min_kb: int = 30, must_have_text: list[str] | None = None) -> list[str]:
    issues = []
    if not p.exists():
        return [f"MISSING: {p.name}"]
    sz = p.stat().st_size
    if sz < min_kb * 1024:
        issues.append(f"SMALL ({sz} B): {p.name}")
    d = fitz.open(p)
    widgets = sum(len(list(pg.widgets() or [])) for pg in d)
    if widgets > 0:
        issues.append(f"NOT FLATTENED ({widgets} widgets): {p.name}")
    text = "\n".join(pg.get_text() for pg in d)
    d.close()
    if must_have_text:
        for needle in must_have_text:
            if needle not in text:
                issues.append(f"MISSING expected text {needle!r}: {p.name}")
    return issues


# ===== Step 13-15: draft email + present + send =====
def assemble_email_body(disciplines_set: frozenset, project_address_short: str,
                        tpr: str, deficiency_date: str) -> tuple[str, str]:
    nice = " + ".join(sorted(disciplines_set, key={"SPK":0,"FA":1,"FP":2}.get))
    subject = f"{project_address_short} — Plan Review Approval Package (DOB Submittal)"
    bullets = [
        "1. Approved Notification of Intent (NOI) — DOB-stamped",
        "2. Plan Review Deficiency Report(s) — initial review + corrections verified",
        "3. Plan Review Approval Certificate and Report (DOB Form) — signed",
        "4. Plan Review Certification Letter — addressed to DOB / Mayda Colon",
        "5. Large Stamp — Construction Codes Verified",
    ]
    body = (f"Hi {{name}},\n\n"
            f"Good news — the third-party plan review for {project_address_short} ({nice}) has been approved. "
            f"The DOB has issued the Notice of Approval Certification Number {tpr} (acceptance date {deficiency_date}).\n\n"
            "Attached are the supporting documents for your DOB permit submittal. "
            "Please submit these alongside your latest drawings:\n\n"
            + "\n".join(bullets) + "\n\n"
            "All PDFs have been flattened (non-editable) per DOB submittal requirements.\n\n"
            "Let me know if anything needs to be revised. Glad to support the next phase — "
            "please reach out when you're ready for the post-permit inspections.")
    return subject, body


# ===== ENTRY POINT =====
def main(urls: list[str], do_send: bool, recipient_override: tuple[str, str] | None) -> None:
    import datetime as dt
    if not urls:
        fail("Pass at least one Google Spreadsheet URL.")

    # Step 1: file IDs
    file_ids = [extract_file_id(u) for u in urls]
    print(f"[1] file_ids: {file_ids}")

    # Step 2: download (Drive MCP cache)
    sheet_paths: list[Path] = []
    cache = REPO / "_drive_cache"
    cache.mkdir(exist_ok=True)
    for fid in file_ids:
        x = cache / f"{fid}.xlsx"
        download_sheet_via_mcp(fid, x)
        sheet_paths.append(x)
    print(f"[2] downloaded: {[p.name for p in sheet_paths]}")

    # Step 3: read metadata
    metas = [read_sheet_metadata(p) for p in sheet_paths]
    print(f"[3] metas: {[(m['discipline_tag'], m['project_address']) for m in metas]}")
    addresses = {m["project_address"] for m in metas}
    if len(addresses) > 1:
        fail(f"Sheets describe different addresses: {addresses}")
    project_address_short = addresses.pop()
    discipline_tags = sorted({m["discipline_tag"] for m in metas})
    if "UNKNOWN" in discipline_tags:
        fail(f"Could not infer discipline for one of the sheets — got: {discipline_tags}")
    disc_set = frozenset(discipline_tags)
    print(f"[3] project={project_address_short!r}  disciplines={disc_set}")

    # Step 4: locate project folder
    proj = locate_project_folder(project_address_short)
    print(f"[4] project_folder: {proj}")
    wrapup_dir = proj / "Wrap up"; wrapup_dir.mkdir(exist_ok=True)

    # Step 5: locate approved NOI
    noi = locate_approved_noi(proj)
    print(f"[5] approved_noi: {noi.name}")

    # Step 6: parse approved NOI
    noi_data = parse_approved_noi(noi)
    print(f"[6] tpr={noi_data['tpr']}  deficiency_date={noi_data['deficiency_date']}")

    # Step 7: today
    today = dt.date.today()
    today_dash = today.strftime("%m-%d-%Y")
    today_slash = today.strftime("%m/%d/%Y")
    day_month = today.strftime("%B ") + f"{today.day}{_ord_suffix(today.day)}"
    year2 = today.strftime("%y")

    # Step 8: discipline copy strings
    disc_blank = DISCIPLINE_BLANK_TABLE.get(disc_set) or "/".join(sorted(disc_set))
    letter_disciplines = LETTER_DISCIPLINES_TABLE.get(disc_set) or list(disc_set)
    largestamp_rows = LARGESTAMP_ROWS_TABLE.get(disc_set) or list(disc_set)
    doc3_disc_rows = DOC3_DISC_ROWS_TABLE.get(disc_set) or list(disc_set)

    # Step 9: recipient
    if recipient_override:
        recipient_name, recipient_email = recipient_override
    else:
        rec = lookup_recipient(project_address_short)
        if not rec:
            fail(f"Recipient email not found in {TRACKER.name} for {project_address_short!r}. "
                 f"Pass --to NAME EMAIL or add row to tracker.")
        recipient_name, recipient_email = rec
    print(f"[9] recipient: {recipient_name} <{recipient_email}>")

    # Step 10: produce 5 docs
    project_address_full = f"{project_address_short}, Washington DC"
    project_address_doc5 = f"{project_address_short}, Washington D.C."
    addr_short_no_punct = project_address_short
    disc_tag = "FP" if disc_set & {"FA", "SPK"} else next(iter(disc_set))

    doc1 = wrapup_dir / noi.name
    shutil.copy2(noi, doc1)
    print(f"  [Doc1] {doc1.name}")

    doc2_pdfs = []
    for meta, sheet_xlsx in zip(metas, sheet_paths):
        out = wrapup_dir / f"{addr_short_no_punct} {meta['discipline_tag']} Plan Review Deficiency Report.pdf"
        print_xlsx_to_pdf(sheet_xlsx, out, meta["filled_last_row"], meta["filled_last_col"])
        doc2_pdfs.append(out)
        print(f"  [Doc2-{meta['discipline_tag']}] {out.name}")

    doc3 = wrapup_dir / f"{addr_short_no_punct} {disc_tag} Plan Review Approval Certificate and Report.pdf"
    fill_doc3(doc3, tpr=noi_data["tpr"], project_name=noi_data["project_name"],
              project_address=project_address_full, deficiency_date=noi_data["deficiency_date"],
              today_date=today_dash, day_month=day_month, year2=year2,
              disciplines=doc3_disc_rows, disc_blank=disc_blank)
    print(f"  [Doc3] {doc3.name}")

    doc4_docx = wrapup_dir / f"{addr_short_no_punct} Plan Review Certification Letter.docx"
    build_doc4(doc4_docx, project_address_full, today_dash, letter_disciplines)
    doc4_pdf = doc4_docx.with_suffix(".pdf")
    docx_to_pdf(doc4_docx, doc4_pdf)
    print(f"  [Doc4] {doc4_pdf.name}")

    doc5_docx = wrapup_dir / f"Large Stamp for {addr_short_no_punct} {disc_tag} review.docx"
    build_doc5(doc5_docx, project_address_doc5, today_slash, largestamp_rows)
    doc5_pdf = doc5_docx.with_suffix(".pdf")
    docx_to_pdf(doc5_docx, doc5_pdf)
    print(f"  [Doc5] {doc5_pdf.name}")

    # Step 11: flatten ALL
    pdfs_to_flatten = [doc1, *doc2_pdfs, doc3, doc4_pdf, doc5_pdf]
    print("[11] flattening …")
    for p in pdfs_to_flatten:
        flatten_pdf(p)

    # Step 12: self-review
    must_in_doc3 = [noi_data["tpr"], today_dash, project_address_short]
    issues = []
    issues += self_review_pdf(doc1, min_kb=200)
    for p in doc2_pdfs: issues += self_review_pdf(p, min_kb=80, must_have_text=[project_address_short])
    issues += self_review_pdf(doc3, min_kb=80, must_have_text=must_in_doc3)
    issues += self_review_pdf(doc4_pdf, min_kb=80, must_have_text=[project_address_short])
    issues += self_review_pdf(doc5_pdf, min_kb=80, must_have_text=[project_address_short])
    if issues:
        print("[12] SELF-REVIEW ISSUES:")
        for x in issues: print(f"    - {x}")
        fail("Halting — fix issues before send.")
    print("[12] self-review: clean")

    # Step 13-14: draft email + preview
    subject, body_template = assemble_email_body(disc_set, project_address_short,
                                                 noi_data["tpr"], noi_data["deficiency_date"])
    body = body_template.replace("{name}", recipient_name.split()[0] if recipient_name and recipient_name != "(see tracker)" else "there")
    PENDING.mkdir(parents=True, exist_ok=True)
    draft = PENDING / f"{addr_short_no_punct.replace(' ','_')}_Wrapup_Draft.md"
    draft.write_text(
        f"**To:** {recipient_name} <{recipient_email}>\n"
        f"**From:** admin@buildingcodeconsulting.com\n"
        f"**CC:** ycao@buildingcodeconsulting.com\n"
        f"**Subject:** {subject}\n\n"
        f"**Attachments ({len(pdfs_to_flatten)}):**\n"
        + "\n".join(f"  - {p.name}  ({p.stat().st_size//1024} KB)" for p in pdfs_to_flatten)
        + f"\n\n---\n\n{body}\n", encoding="utf-8")
    print(f"\n[13] draft: {draft}")
    print(f"[14] preview ↓\n")
    print(draft.read_text(encoding="utf-8"))

    # Step 15: send gate
    if not do_send:
        print(f"\n--- Email content ready for {recipient_name}. Reply 'Y' to send "
              f"(re-run with --send when approved). ---")
        return

    # Step 16: SMTP
    from email_sender import send_from_admin_with_attachments
    print(f"[16] sending → {recipient_email}")
    success, info = send_from_admin_with_attachments(
        to_email=recipient_email, subject=subject, body_plain=body,
        attachment_paths=[str(p) for p in pdfs_to_flatten],
    )
    print(f"     result: success={success}  info={info}")
    if success:
        sent = draft.with_name(draft.stem + "-SENT.md"); draft.rename(sent)
        print(f"     audit: {sent.name}")


def _ord_suffix(n: int) -> str:
    return "th" if 10 <= n % 100 <= 20 else {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("urls", nargs="+", help="Google Spreadsheet URLs (one per discipline)")
    ap.add_argument("--send", action="store_true", help="Actually send after self-review (still requires preflight pass)")
    ap.add_argument("--to", nargs=2, metavar=("NAME", "EMAIL"), help="Override recipient")
    args = ap.parse_args()
    main(args.urls, do_send=args.send, recipient_override=tuple(args.to) if args.to else None)
