"""Fix resume:
- Restore the original verbose 'Building and Electrical Plan Review' description
  (I shouldn't have trimmed it).
- Add ECS Mid-Atlantic LLC and dmy capitol llc as employer entries
  (between Frederick County and YIDA in document order).
- Keep BCC + APRG entries already added at top.
- Keep PE Fire Protection cert update.
"""
from pathlib import Path
from docx import Document

SRC = Path(r"C:\Users\Kyle Cao\DC Business\Legal Docs\Kyle Cao's Resume.docx")
doc = Document(str(SRC))

ORIGINAL_DESC = (
    "Building, electrical, fire safety inspections and plan review for all kinds "
    "of building projects. Including but not limit to, mixed use apartments buildings, "
    "hotels, health care facilities, theaters, museums, zero energy green buildings, "
    "university power plant and facility buildings, office buildings, classrooms, "
    "recreation center, football stadium and facility buildings, elementary school, "
    "high school new building and existing renovations and residential buildings and "
    "Medical Center etc."
)

# 1) Restore the trimmed description
for p in doc.paragraphs:
    if p.text.startswith("Plan Review and Code Compliance Inspection for all types") and \
       p.runs and p.runs[0].italic:
        # Find which entry — must be the one BETWEEN "DC, VA and MD areas..." and "Micron..."
        # Identify by neighbors:
        prev = p._p.getprevious()
        prev_text = prev.text if prev is not None else ""
        if "DC, VA and MD" in prev_text:
            for r in list(p.runs):
                r.text = ""
            p.runs[0].text = ORIGINAL_DESC
            p.runs[0].italic = True
            break

# 2) Insert dmy + ECS entries before YIDA
yida = None
for p in doc.paragraphs:
    if p.text.startswith("Structural Engineer of the 5") or "YIDA" in p.text:
        yida = p
        break

style = yida.style


def insert_entry(ref, title, loc, date, desc):
    p1 = ref.insert_paragraph_before(style=style)
    r1 = p1.add_run(title)
    r1.bold = True
    p2 = ref.insert_paragraph_before(style=style)
    p2.add_run(loc + "\t")
    r2 = p2.add_run(date)
    r2.italic = True
    p3 = ref.insert_paragraph_before(style=style)
    r3 = p3.add_run(desc)
    r3.italic = True


# Insert in document order so resulting order is dmy → ECS → YIDA
insert_entry(
    yida,
    "dmy capitol llc",
    "Chantilly, VA",
    "Dec. 2014 – Mar. 2018",
    "Architectural MEP design and fire protection system design for residential and "
    "commercial projects as Project Engineer.",
)
insert_entry(
    yida,
    "ECS Mid-Atlantic LLC",
    "Chantilly, VA",
    "Jul. 2014 – Nov. 2014",
    "Civil and geotechnical engineering project work.",
)

doc.save(str(SRC))

print(f"Saved: {SRC}\n")
print("=== Final resume ===")
for i, p in enumerate(Document(str(SRC)).paragraphs):
    print(f"{i:3d} | {p.text[:130]}")
