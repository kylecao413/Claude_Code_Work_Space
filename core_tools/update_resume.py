"""Update Yue Cao's resume: add BCC + American Plan Review Group entries
at top of RELEVANT EXPERIENCE; add PE Fire Protection to PE list; remove
'Candidate of - Fire Protection PE' line.
"""
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

SRC = Path(r"C:\Users\Kyle Cao\DC Business\Legal Docs\Kyle Cao's Resume.docx")
OUT = SRC  # overwrite in place

doc = Document(str(SRC))
paras = doc.paragraphs


def find_idx(prefix: str) -> int:
    for i, p in enumerate(paras):
        if p.text.startswith(prefix):
            return i
    raise SystemExit(f"not found: {prefix}")


# anchor = first existing experience entry; we'll insert NEW entries before it
anchor = paras[find_idx("Building and Electrical Plan Review and Inspection")]
style = anchor.style


def insert_entry(ref, title, loc, date, desc):
    p1 = ref.insert_paragraph_before(style=style)
    r1 = p1.add_run(title)
    r1.bold = True

    p2 = ref.insert_paragraph_before(style=style)
    p2.add_run(loc + "\t")
    r2b = p2.add_run(date)
    r2b.italic = True

    p3 = ref.insert_paragraph_before(style=style)
    r3 = p3.add_run(desc)
    r3.italic = True


# Order: most recent first → BCC, then American Plan Review Group
insert_entry(
    anchor,
    "Building Code Consulting LLC",
    "Washington, DC",
    "Aug. 2023 – until now",
    "Plan Review and Code Compliance Inspection for all types of building projects.",
)
insert_entry(
    anchor,
    "American Plan Review Group",
    "Austin, TX (remote)",
    "2024 – until now",
    "Plan Review and Code Compliance Inspection for all types of building projects.",
)

# refresh paras list after insertions
paras = doc.paragraphs

# === Trim the verbose description on the existing 'Building and Electrical Plan Review' entry
# to keep resume within 1 page. Original is ~5 lines of building-type enumeration.
for i, p in enumerate(paras):
    if p.text.startswith("Building, electrical, fire safety inspections and plan review for all kinds"):
        # wipe runs and replace with a single concise italic line
        for r in list(p.runs):
            r.text = ""
        p.runs[0].text = (
            "Plan Review and Code Compliance Inspection for all types of building projects."
        )
        p.runs[0].italic = True
        break

# === Update CERTIFICATE & LICENSE block ===
# 1) Add "- PE Fire Protection" after "- PE Civil – Geotechnical"
geo_idx = find_idx_text = None
for i, p in enumerate(paras):
    if "PE Civil" in p.text and "Geotechnical" in p.text:
        geo_idx = i
        break
geo_para = paras[geo_idx]
new_fp = geo_para._p.makeelement(qn("w:p"), {})
geo_para._p.addnext(new_fp)
# Build fresh paragraph by inserting before the next sibling
new_para_obj = doc.paragraphs[geo_idx + 1]  # the one we just inserted
# our inserted bare <w:p> has no style/runs — replace with proper insert_paragraph_before
# easier: drop that bare element and use insert_paragraph_before on the next non-empty para
geo_para._p.getparent().remove(new_fp)
# refresh
paras = doc.paragraphs
# the paragraph AFTER geo_idx is "Licensed Master of Electrician..." → insert before it
master_idx = None
for i, p in enumerate(paras):
    if p.text.startswith("Licensed Master of Electrician"):
        master_idx = i
        break
fp_para = paras[master_idx].insert_paragraph_before(style=paras[master_idx].style)
fp_para.add_run("               - PE Fire Protection")

# 2) Remove "Candidate of - Fire Protection PE" line
paras = doc.paragraphs
for p in paras:
    if p.text.strip().startswith("Candidate of") and "Fire Protection" in p.text:
        p._p.getparent().remove(p._p)
        break

doc.save(str(OUT))
print(f"Saved: {OUT}")
print("\n=== Resume after update ===")
for i, p in enumerate(Document(str(OUT)).paragraphs):
    print(f"{i:3d} | {p.text[:130]}")
