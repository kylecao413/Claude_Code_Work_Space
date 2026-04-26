"""One-shot: generate Rexfield 1522 Rhode Island Ave NE plan review proposal from 1812 H Pl template."""
import shutil
from pathlib import Path
from docx import Document

SRC = Path(r"C:\Users\Kyle Cao\DC Business\Building Code Consulting\Projects\Nery Soto Fire Protection\1812 H PL NE - PR\1812 - 1820 H Pl NE Plan Review Proposal.docx")
DST_DIR = Path(r"C:\Users\Kyle Cao\DC Business\Building Code Consulting\Projects\Nery Soto Fire Protection\1522 Rhode Island Ave NE - PR")
DST = DST_DIR / "1522 Rhode Island Ave NE Plan Review Proposal.docx"

NEW_PROJECT_TITLE = "1522 Rhode Island Ave NE"
NEW_DATE = "April 17, 2026"
NEW_DESCRIPTION = (
    "The project consists of the review of fire alarm and fire sprinkler construction "
    "documents for a new five-story plus penthouse residential apartment building located "
    "at 1522 Rhode Island Ave NE in Washington, D.C. Building Code Consulting will serve "
    "as the third-party plan reviewer, evaluating the submitted fire protection drawings "
    "for compliance with the District of Columbia Building Code, the DC Fire Code, and "
    "applicable NFPA standards (including NFPA 13 sprinkler design with multiple hydraulic "
    "calculation areas and NFPA 72 fire alarm system design with battery calculations), "
    "and advising the client on all matters related to code conformity prior to AHJ submission."
)
NEW_FEE = "$2,500.00"


def replace_para_preserving_format(p, new_text: str):
    """Replace paragraph text while preserving the font/size/bold of the first run."""
    if p.runs:
        first = p.runs[0]
        for r in p.runs[1:]:
            r.text = ""
        first.text = new_text
    else:
        p.add_run(new_text)


def main():
    DST_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SRC, DST)

    d = Document(DST)

    # Cover page title (para 3)
    replace_para_preserving_format(d.paragraphs[3], NEW_PROJECT_TITLE)
    # Date (para 12)
    replace_para_preserving_format(d.paragraphs[12], NEW_DATE)
    # Project Description (para 83)
    replace_para_preserving_format(d.paragraphs[83], NEW_DESCRIPTION)

    # Estimated Fee table (Table 0, row 0, cell 1)
    fee_cell = d.tables[0].rows[0].cells[1]
    for p in fee_cell.paragraphs:
        if "$" in p.text:
            replace_para_preserving_format(p, NEW_FEE)
            break

    d.save(DST)

    # Verify
    d2 = Document(DST)
    print(f"Saved: {DST}")
    print(f"  Title:        {d2.paragraphs[3].text}")
    print(f"  Date:         {d2.paragraphs[12].text}")
    print(f"  Client:       {d2.paragraphs[7].text}")
    print(f"  Contact:      {d2.paragraphs[10].text}")
    print(f"  Description:  {d2.paragraphs[83].text[:120]}...")
    print(f"  Fee:          {d2.tables[0].rows[0].cells[1].text}")


if __name__ == "__main__":
    main()
