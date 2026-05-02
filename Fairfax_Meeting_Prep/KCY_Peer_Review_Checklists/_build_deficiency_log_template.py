"""Generate KCY Engineering Code Consulting LLC peer-review deficiency log template (.docx)."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_background(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tbl_pr.append(borders)


doc = Document()

# Page margins
for section in doc.sections:
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

# ---------- Letterhead block ----------
header_p = doc.add_paragraph()
header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = header_p.add_run("KCY ENGINEERING CODE CONSULTING LLC")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = sub_p.add_run("Third-Party Plan Peer Review")
sub.italic = True
sub.font.size = Pt(11)

addr_p = doc.add_paragraph()
addr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
addr = addr_p.add_run("[Street Address] | [City, State ZIP] | [Phone] | [Email] | [Website]")
addr.font.size = Pt(9)
addr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# Horizontal rule (using bottom border on a paragraph)
hr_p = doc.add_paragraph()
hr_pPr = hr_p._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "8")
bottom.set(qn("w:space"), "1")
bottom.set(qn("w:color"), "003366")
pBdr.append(bottom)
hr_pPr.append(pBdr)

# ---------- Title ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("PLAN REVIEW DEFICIENCY LOG")
tr.bold = True
tr.font.size = Pt(14)

# ---------- Project info block ----------
info_table = doc.add_table(rows=4, cols=4)
info_table.autofit = False
info_table.columns[0].width = Inches(1.4)
info_table.columns[1].width = Inches(2.6)
info_table.columns[2].width = Inches(1.4)
info_table.columns[3].width = Inches(1.7)
set_table_borders(info_table)

info_rows = [
    ("Project Name:", "[Project Name]", "Permit Number:", "[Permit #]"),
    ("Project Address:", "[Street Address, City, State ZIP]", "Parcel ID:", "[Parcel ID]"),
    ("Discipline(s):", "[Building / Mechanical / Electrical / Plumbing / Fire Protection / Structural / Gas]", "Review Date:", "[YYYY-MM-DD]"),
    ("Designer of Record:", "[Name, License #]", "Submission #:", "[1st / 2nd / Re-review]"),
]
for r, (a, b, c, d) in enumerate(info_rows):
    cells = info_table.rows[r].cells
    for cell, text, bold in zip(cells, (a, b, c, d), (True, False, True, False)):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(text)
        run.font.size = Pt(10)
        run.bold = bold
    # shade label cells
    set_cell_background(cells[0], "EAF1F8")
    set_cell_background(cells[2], "EAF1F8")

doc.add_paragraph()

# ---------- Reviewer info ----------
rev_table = doc.add_table(rows=2, cols=4)
rev_table.autofit = False
rev_table.columns[0].width = Inches(1.4)
rev_table.columns[1].width = Inches(2.6)
rev_table.columns[2].width = Inches(1.4)
rev_table.columns[3].width = Inches(1.7)
set_table_borders(rev_table)

rev_rows = [
    ("Peer Reviewer:", "[Yue Cao, PE, MCP] — KCY Engineering Code Consulting LLC", "PR No.:", "[PR Number]"),
    ("Code Edition:", "2021 VCC / 2021 VEBC / 2021 IMC / 2021 IPC / 2020 NEC / ICC A117.1-2009", "Page:", "1 of __"),
]
for r, (a, b, c, d) in enumerate(rev_rows):
    cells = rev_table.rows[r].cells
    for cell, text, bold in zip(cells, (a, b, c, d), (True, False, True, False)):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(text)
        run.font.size = Pt(10)
        run.bold = bold
    set_cell_background(cells[0], "EAF1F8")
    set_cell_background(cells[2], "EAF1F8")

doc.add_paragraph()

# ---------- Scope statement ----------
scope_p = doc.add_paragraph()
sr = scope_p.add_run(
    "Scope. KCY Engineering Code Consulting LLC has performed a third-party peer review "
    "of the construction documents listed above against the applicable code editions. "
    "The deficiencies listed below are observations of non-compliance identified on the "
    "drawings and supporting documents. Code citations are provided for reference. "
    "The design professional of record retains sole responsibility for the design and "
    "for determining the appropriate corrective action."
)
sr.font.size = Pt(9)
sr.italic = True

doc.add_paragraph()

# ---------- Deficiency table ----------
heading = doc.add_paragraph()
hr = heading.add_run("Deficiencies")
hr.bold = True
hr.font.size = Pt(12)

# Columns: # | Discipline | Sheet/Detail | Deficiency Observation + Code Citation | Status
table = doc.add_table(rows=1, cols=5)
table.autofit = False
table.columns[0].width = Inches(0.4)
table.columns[1].width = Inches(0.9)
table.columns[2].width = Inches(1.3)
table.columns[3].width = Inches(3.5)
table.columns[4].width = Inches(1.0)
set_table_borders(table)

headers = ["#", "Discipline", "Sheet / Detail Location", "Deficiency Observation + Code Citation", "Status"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(h)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_background(cell, "003366")

# Example placeholder rows
example_rows = [
    ("1", "Building", "A0.0 / Coversheet",
     "Building Plan Review Cover Sheet accessibility cost block is not completed; total applicable cost of construction, items being upgraded, and associated upgrade costs are not provided. (VEBC Chapter 4)",
     "Open"),
    ("2", "Building", "A1.1 / Life Safety Plan",
     "Door schedule and door hardware schedule are not provided; door operation compliance with VCC §1010.1.9 cannot be determined. (VCC §109.4, §1010.1.9)",
     "Open"),
    ("3", "Structural", "S0.0 / General Notes",
     "Statement of Special Inspections is missing the Structural Engineer's seal and signature; pages 2–3 contain items unmarked. (VCC §109.4, Chapter 17)",
     "Open"),
]
for nm, disc, sheet, obs, status in example_rows:
    row = table.add_row().cells
    for i, val in enumerate((nm, disc, sheet, obs, status)):
        row[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = row[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(val)
        run.font.size = Pt(9)
        if i == 0 or i == 4:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add 12 blank rows for the user to fill
for _ in range(12):
    row = table.add_row().cells
    for i in range(5):
        p = row[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("")
        run.font.size = Pt(9)

# ---------- Footer / status legend ----------
doc.add_paragraph()
legend = doc.add_paragraph()
lr = legend.add_run("Status legend: Open = unresolved · Resolved = corrected on resubmission · Withdrawn = not a deficiency on further review")
lr.font.size = Pt(8)
lr.italic = True
lr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

discl = doc.add_paragraph()
dr = discl.add_run(
    "This deficiency log identifies observations of code non-compliance. "
    "Recommended corrections are not provided herein; the design professional of record is responsible for "
    "determining and implementing appropriate corrective measures and for resolving any code interpretation. "
    "This document does not waive any code requirement or any obligation under the building official's review."
)
dr.font.size = Pt(8)
dr.italic = True
dr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# Signature block
doc.add_paragraph()
sig = doc.add_paragraph()
sigr = sig.add_run("Reviewed and signed:")
sigr.bold = True
sigr.font.size = Pt(10)

sig_table = doc.add_table(rows=2, cols=2)
sig_table.autofit = False
sig_table.columns[0].width = Inches(3.5)
sig_table.columns[1].width = Inches(3.5)
for r, (a, b) in enumerate([
    ("_______________________________________", "_______________________________________"),
    ("Yue Cao, PE, MCP — Peer Reviewer (PR #)", "Date"),
]):
    cells = sig_table.rows[r].cells
    for c, text in zip(cells, (a, b)):
        p = c.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9)
        if r == 1:
            run.italic = True

out = Path(__file__).parent / "00_KCY_Deficiency_Log_Template.docx"
doc.save(str(out))
print(f"Saved: {out}")
