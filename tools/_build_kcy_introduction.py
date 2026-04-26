"""
_build_kcy_introduction.py — One-shot builder for
`../Marketing/KCY Engineer PLLC Introduction.docx`.

Mirrors the structure of the BCC "Building Code Consulting LLC Introduction.pdf.docx"
(title at 18pt bold; bold H2 section headers; plain-run bullets with a bold lead-in
phrase) but swaps in KCY brand + services + territory per Kyle's 2026-04-23 spec.

Run once:
    python tools/_build_kcy_introduction.py

Safe to re-run — overwrites the output file.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt


OUT_PATH = (
    Path(__file__).resolve().parent.parent.parent
    / "Marketing" / "KCY Engineer PLLC Introduction.docx"
)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def add_section_header(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True


def add_body(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_lead_in_bullet(doc: Document, lead: str, rest: str) -> None:
    """Paragraph whose first phrase is bold (lead), rest is plain."""
    p = doc.add_paragraph()
    r1 = p.add_run(lead)
    r1.bold = True
    p.add_run(rest)


def add_blank(doc: Document) -> None:
    doc.add_paragraph()


def build() -> None:
    doc = Document()

    # ── Title ──
    add_title(doc, "KCY Engineer PLLC")
    add_blank(doc)

    # ── About Us ──
    add_section_header(doc, "About Us")
    add_body(doc, (
        "KCY Engineer PLLC is a licensed Professional Engineering firm providing "
        "Expedited Peer Review services nationwide, along with Third-Party Code "
        "Compliance Inspections in select Virginia and Maryland jurisdictions. "
        "Our practice covers all major disciplines — Building, Mechanical, "
        "Electrical, Plumbing, and Fire Protection (BMEPF)."
    ))
    add_body(doc, (
        "The firm is led by a licensed Professional Engineer (PE) who is also "
        "an ICC-certified Master Code Professional (MCP), supported by "
        "collaborating licensed PEs across each of the BMEPF disciplines. This "
        "combination of engineering seal authority and code-compliance depth "
        "allows KCY to deliver reviews and inspections that are both technically "
        "rigorous and responsive to construction schedules."
    ))
    add_body(doc, (
        "We approach every engagement as a hands-on technical partner — helping "
        "clients resolve code compliance questions efficiently, minimize "
        "revision cycles, and keep projects moving."
    ))
    add_blank(doc)

    # ── Why Choose Us ──
    add_section_header(doc, "Why Choose Us")
    add_lead_in_bullet(doc,
        "Multi-Discipline PE Expertise: ",
        "Licensed Professional Engineers across Building, Mechanical, "
        "Electrical, Plumbing, and Fire Protection disciplines — with ICC "
        "Master Code Professional (MCP) certification on the core team."
    )
    add_lead_in_bullet(doc,
        "Expedited Peer Review — Nationwide: ",
        "We accept Expedited Peer Review engagements in any U.S. jurisdiction, "
        "with turnaround typically measured in days rather than weeks."
    )
    add_lead_in_bullet(doc,
        "Responsive Inspection Scheduling: ",
        "Same-day or next-business-day inspection availability in our served "
        "inspection territories, so construction milestones are not delayed "
        "by inspector access."
    )
    add_lead_in_bullet(doc,
        "Visit-Based Billing: ",
        "Inspection billing is a flat rate per visit actually performed — "
        "clients are never billed against an upfront estimate. If a project "
        "wraps in fewer inspections than projected, clients pay only for what "
        "was done."
    )
    add_lead_in_bullet(doc,
        "Clear Engineering Voice: ",
        "Reports are written by the engineer of record. Findings are "
        "specific, actionable, and backed by code citation — suitable for "
        "submission to the AHJ without translation or rework."
    )
    add_blank(doc)

    # ── Our Services ──
    add_section_header(doc, "Our Services")
    add_lead_in_bullet(doc,
        "Expedited Peer Review (Nationwide): ",
        "Independent third-party peer review of construction documents "
        "across all BMEPF disciplines. Identifies code deficiencies before "
        "AHJ submission — reducing agency review cycles and avoiding costly "
        "revision loops."
    )
    add_lead_in_bullet(doc,
        "Third-Party Code Compliance Inspections — Virginia (Residential): ",
        "Northern Virginia residential third-party inspection services, "
        "currently active in Fairfax County. Additional NoVA jurisdictions "
        "(Arlington, Alexandria, Loudoun, Prince William, Stafford) pending "
        "AHJ registration."
    )
    add_lead_in_bullet(doc,
        "Third-Party Code Compliance Inspections — Prince George's County, MD: ",
        "Full commercial and residential third-party inspection services in "
        "Prince George's County, Maryland — BMEPF disciplines."
    )
    add_lead_in_bullet(doc,
        "Code Consulting: ",
        "Engineering-led code consulting to help project teams navigate "
        "complex code-interpretation questions, conflicting requirements, "
        "and variance/modification strategy — throughout Maryland, Virginia, "
        "and the broader Mid-Atlantic region."
    )
    add_blank(doc)

    # ── Service Commitments ──
    add_section_header(doc, "Service Commitments")
    add_lead_in_bullet(doc,
        "Timely Turnaround: ",
        "Peer review engagements are returned on a defined, expedited schedule "
        "agreed with the client at project start. Inspection requests are "
        "accommodated within one business day."
    )
    add_lead_in_bullet(doc,
        "Clear Communication: ",
        "Detailed reports delivered directly by the reviewing / inspecting "
        "engineer. No hand-offs to non-engineering staff for technical questions."
    )
    add_lead_in_bullet(doc,
        "Technical Accuracy: ",
        "All reviews and inspections are conducted or supervised by licensed "
        "Professional Engineers in the applicable discipline."
    )
    add_lead_in_bullet(doc,
        "Reliable Support: ",
        "Available weekdays with flexibility for schedule-critical project demands."
    )
    add_blank(doc)

    # ── Contact Us ──
    add_section_header(doc, "Contact Us")
    p = doc.add_paragraph()
    r = p.add_run("KCY Engineer PLLC")
    r.bold = True
    p.add_run("\nProfessional In Charge: Kyle Cao, PE, MCP")
    p.add_run("\nPhone: (571) 365-6937")
    p.add_run("\nEmail: ycao@kcyengineer.com")
    p.add_run("\nWebsite: www.kcyengineer.com")

    # Save
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Wrote: {OUT_PATH}")


if __name__ == "__main__":
    build()
